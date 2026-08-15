"""
Table S5 (Appendix): Post hoc formal minimum sample-size justification by
cohort (Riley et al. 2020, BMJ 368:m441), read directly from
outputs/pmsampsize_results.xlsx (src/22_pmsampsize.py) - not retyped.

Landscape (11 data columns). Three-line rule style (rule above header, below
header, at foot; no vertical rules, no grid) matching the main-manuscript
Table 1/2 convention (src/25_tables_1_2.py), Times New Roman. "Adequate?"
bolded per row so the flare-vs-ESRD split is visible at a glance.

Saves: outputs/Table_S5_SampleSize.docx
"""
import pandas as pd
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT

BASE = "/Users/elizabetharmstrong/Library/CloudStorage/OneDrive-ImperialCollegeLondon/Lupus Project"
OUT = f"{BASE}/outputs"
OUTPUT = f"{OUT}/Table_S5_SampleSize.docx"
FONT = "Times New Roman"

df = pd.read_excel(f"{OUT}/pmsampsize_results.xlsx")
print(df.to_string(index=False))


def set_cell_borders(cell, top=None, bottom=None, left=None, right=None, insideH=None, insideV=None):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    specs = {"top": top, "bottom": bottom, "left": left, "right": right,
             "insideH": insideH, "insideV": insideV}
    for edge, spec in specs.items():
        el = tcBorders.find(qn(f'w:{edge}'))
        if el is None:
            el = OxmlElement(f'w:{edge}')
            tcBorders.append(el)
        if spec is None:
            el.set(qn('w:val'), 'nil')
        else:
            el.set(qn('w:val'), 'single')
            el.set(qn('w:sz'), str(spec))
            el.set(qn('w:color'), '000000')
            el.set(qn('w:space'), '0')


def clear_all_borders(cell):
    set_cell_borders(cell)


def apply_three_line_borders(table, header_row, last_row, n_cols):
    THICK, THIN = 18, 8
    # When header_row == 0, the top-rule and header-bottom-rule cells are the
    # SAME cell - set_cell_borders() rebuilds all 6 edges every call, so two
    # separate calls would let the second silently overwrite the first's top
    # edge back to nil. Combine into one call whenever they coincide (caught
    # via docx readback verification: row-0 top was rendering as 'nil').
    if header_row == 0:
        for c in range(n_cols):
            set_cell_borders(table.cell(0, c), top=THICK, bottom=THIN)
    else:
        for c in range(n_cols):
            set_cell_borders(table.cell(0, c), top=THICK)
        for c in range(n_cols):
            set_cell_borders(table.cell(header_row, c), bottom=THIN)
    for c in range(n_cols):
        set_cell_borders(table.cell(last_row, c), bottom=THICK)


def set_table_cell_margins(table, top_pt=3, bottom_pt=3, left_pt=5, right_pt=5):
    tbl = table._tbl
    tblPr = tbl.tblPr
    cellMar = OxmlElement('w:tblCellMar')
    for edge, pt_val in (('top', top_pt), ('bottom', bottom_pt), ('left', left_pt), ('right', right_pt)):
        node = OxmlElement(f'w:{edge}')
        node.set(qn('w:w'), str(int(pt_val * 20)))
        node.set(qn('w:type'), 'dxa')
        cellMar.append(node)
    tblPr.append(cellMar)


def set_run_font(run, size=10, bold=False, italic=False):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
        rFonts.set(qn(attr), FONT)


def set_cell(cell, text, size=10, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ""
    p = cell.paragraphs[0]
    for r in list(p.runs):  # cell.text = "" can leave a stray unformatted empty run
        r._element.getparent().remove(r._element)
    p.alignment = align
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic)


doc = Document()
section = doc.sections[0]
section.orientation = WD_ORIENT.LANDSCAPE
section.page_width, section.page_height = section.page_height, section.page_width
section.top_margin = section.bottom_margin = Cm(2.5)
section.left_margin = section.right_margin = Cm(2.5)

title = doc.add_paragraph()
run = title.add_run(
    "Table S5. Post hoc formal minimum sample-size justification by cohort "
    "(Riley et al. 2020, BMJ 368:m441)."
)
set_run_font(run, size=11, bold=True)
title.paragraph_format.space_after = Pt(8)

headers = ["Cohort", "p", "Prevalence", "C-stat\n(LR, BC)", "Cox-Snell\nR²", "n\nRequired",
           "Events\nReq.", "EPP\nReq.", "n\nActual", "Events\nActual", "EPV\nActual", "Adequate?"]
col_widths = [2.6, 1.6, 1.7, 2.1, 1.7, 1.7, 1.8, 1.7, 1.7, 1.8, 1.7, 1.8]

n_rows = 1 + len(df)
table = doc.add_table(rows=n_rows, cols=len(headers))
table.style = "Table Grid"
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False
set_table_cell_margins(table, top_pt=3, bottom_pt=3, left_pt=5, right_pt=5)

for i, w in enumerate(col_widths):
    table.columns[i].width = Cm(w)
for row in table.rows:
    for cell, w in zip(row.cells, col_widths):
        cell.width = Cm(w)

for c, h in enumerate(headers):
    set_cell(table.cell(0, c), h, size=9.5, bold=True)

for r, row in enumerate(df.itertuples(index=False), start=1):
    adequate = row._11  # "Adequate per pmsampsize?"
    set_cell(table.cell(r, 0), row.Cohort, align=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell(table.cell(r, 1), str(row._1))                          # Parameters (p)
    set_cell(table.cell(r, 2), f"{row.Prevalence:.3f}")              # Prevalence
    set_cell(table.cell(r, 3), f"{row._3:.3f}")                      # Anticipated C-statistic
    set_cell(table.cell(r, 4), f"{row._4:.4f}")                      # Derived Cox-Snell R2
    set_cell(table.cell(r, 5), str(row._5))                          # n required
    set_cell(table.cell(r, 6), f"{row._6:.1f}")                      # Events required
    set_cell(table.cell(r, 7), f"{row._7:.2f}")                      # EPP required
    set_cell(table.cell(r, 8), str(row._8))                          # n actual
    set_cell(table.cell(r, 9), str(row._9))                          # Events actual
    set_cell(table.cell(r, 10), f"{row._10:.2f}")                    # EPV actual
    set_cell(table.cell(r, 11), adequate, bold=True)                 # Adequate?

for row in table.rows:
    for cell in row.cells:
        clear_all_borders(cell)
apply_three_line_borders(table, header_row=0, last_row=n_rows - 1, n_cols=len(headers))

footnote = doc.add_paragraph()
run = footnote.add_run(
    "p = number of final predictor parameters. Anticipated C-statistic = the cohort's own "
    "optimism-adjusted (bias-corrected, Harrell bootstrap) Logistic Regression AUROC (the paper's "
    "designated primary model), per Riley et al.'s recommendation to use an optimism-adjusted "
    "C-statistic. Required n = max of three criteria: (1) predictor-effect shrinkage ≤10% loss, "
    "(2) ≤0.05 absolute difference between apparent and adjusted Nagelkerke R², (3) intercept "
    "estimated to within a 0.05 margin of error. All three flare cohorts are underpowered by these "
    "criteria despite meeting the conventional EPV-10 heuristic; both ESRD cohorts are adequately "
    "powered by both criteria."
)
set_run_font(run, size=9, italic=True)
footnote.paragraph_format.space_before = Pt(6)

doc.save(OUTPUT)
print(f"\nSaved: {OUTPUT}")
