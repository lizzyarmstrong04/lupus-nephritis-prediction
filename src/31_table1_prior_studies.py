"""
Table 1. Prior machine-learning studies in lupus nephritis.

This is a literature-review table (citations to prior published studies),
not derived from this project's own data - content is transcribed directly
from the source table the user provided, not independently verified against
external literature (nothing in this repo to check citations against).
Double-check the cited figures/citations against the original papers before
submission.

NOTE: this table is also numbered "Table 1", the same number already used
for the main results table (src/25_tables_1_2.py, "Table 1. Model
performance by cohort..."). That's expected if this table lives in an
earlier section (e.g. Introduction/Background) with its own numbering
sequence, but flagging in case it's an unintended collision - only one
"Table 1" should exist per numbering sequence in the final manuscript.

Three-line rule style (rule above header, below header, at foot; no
vertical rules, no grid), matching the main-manuscript Table 1/2 convention
(src/25_tables_1_2.py), Times New Roman.

Saves: outputs/Table1_Prior_Studies.docx
"""
from docx import Document
from docx.shared import Pt, Cm, Mm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

BASE = "/Users/elizabetharmstrong/Library/CloudStorage/OneDrive-ImperialCollegeLondon/Lupus Project"
OUTPUT = f"{BASE}/outputs/Table1_Prior_Studies.docx"
FONT = "Times New Roman"

HEADERS = ["Study", "Design / cohort", "Method", "Outcome", "Performance", "Key limitation"]
ROWS = [
    ["Chen et al.\n(2021)", "1,694 biopsy-proven patients; 59 variables", "XGBoost",
     "5-year flare", "AUROC 0.819",
     "Long horizon; histopathology entered as composite scores"],
    ["Huang et al.\n(2024)", "Longitudinal time-series (repeat measurements)", "LSTM",
     "Flare", "C-index 0.897",
     "Interpretability limited to global feature importance"],
    ["Chen et al.\n(2022)", "Incorporates polygenic risk scores", "ML + genomic data",
     "Flare", "Improved over baseline (Δ not reported)",
     "Genomic data requirement limits routine applicability"],
]
COL_WIDTHS_CM = [2.5, 3.2, 1.9, 1.8, 2.4, 5.2]  # sums to 17.0 cm


# --- Border / font helpers (matching src/25_tables_1_2.py exactly) ---

def set_cell_borders(cell, top=None, bottom=None, left=None, right=None, insideH=None, insideV=None):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    specs = {"top": top, "bottom": bottom, "left": left, "right": right,
             "insideH": insideH, "insideV": insideV}
    for edge, spec in specs.items():
        el = tcBorders.find(qn(f'w:{edge}'))
        if el is None:
            el = OxmlElement(f'w:{edge}')
            tcBorders.append(el)
        if spec is None:
            el.set(qn('w:val'), 'nil')
        else:
            el.set(qn('w:val'), 'single')
            el.set(qn('w:sz'), str(spec))
            el.set(qn('w:color'), '000000')
            el.set(qn('w:space'), '0')


def clear_all_borders(cell):
    set_cell_borders(cell)


def apply_three_line_borders(table, header_row, last_row, n_cols):
    THICK, THIN = 18, 8
    if header_row == 0:
        for c in range(n_cols):
            set_cell_borders(table.cell(0, c), top=THICK, bottom=THIN)
    else:
        for c in range(n_cols):
            set_cell_borders(table.cell(0, c), top=THICK)
        for c in range(n_cols):
            set_cell_borders(table.cell(header_row, c), bottom=THIN)
    for c in range(n_cols):
        set_cell_borders(table.cell(last_row, c), bottom=THICK)


def set_table_cell_margins(table, top_pt=3, bottom_pt=3, left_pt=5, right_pt=5):
    tbl = table._tbl
    tblPr = tbl.tblPr
    cellMar = OxmlElement('w:tblCellMar')
    for edge, pt_val in (('top', top_pt), ('bottom', bottom_pt), ('left', left_pt), ('right', right_pt)):
        node = OxmlElement(f'w:{edge}')
        node.set(qn('w:w'), str(int(pt_val * 20)))
        node.set(qn('w:type'), 'dxa')
        cellMar.append(node)
    tblPr.append(cellMar)


def set_run_font(run, size=10, bold=False, italic=False):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
        rFonts.set(qn(attr), FONT)


def set_cell(cell, text, size=10, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    for r in list(p.runs):  # cell.text = "" can leave a stray unformatted empty run
        r._element.getparent().remove(r._element)
    p.alignment = align
    lines = text.split("\n")
    for i, line in enumerate(lines):
        if i > 0:
            run = p.add_run()
            run.add_break()
        run = p.add_run(line)
        set_run_font(run, size=size, bold=bold, italic=italic)


doc = Document()
section = doc.sections[0]
section.page_width = Mm(210)   # A4 portrait
section.page_height = Mm(297)
section.top_margin = section.bottom_margin = Cm(2)
section.left_margin = section.right_margin = Cm(2)

title = doc.add_paragraph()
run = title.add_run("Table 1. Prior machine-learning studies in lupus nephritis.")
set_run_font(run, size=11, bold=True)
title.paragraph_format.space_after = Pt(8)

n_rows = 1 + len(ROWS)
n_cols = len(HEADERS)
table = doc.add_table(rows=n_rows, cols=n_cols)
table.style = "Table Grid"
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False
set_table_cell_margins(table, top_pt=3, bottom_pt=3, left_pt=5, right_pt=5)

for i, w in enumerate(COL_WIDTHS_CM):
    table.columns[i].width = Cm(w)
for row in table.rows:
    for cell, w in zip(row.cells, COL_WIDTHS_CM):
        cell.width = Cm(w)

for c, h in enumerate(HEADERS):
    set_cell(table.cell(0, c), h, bold=True)

for r, row_data in enumerate(ROWS, start=1):
    for c, val in enumerate(row_data):
        set_cell(table.cell(r, c), val)

for row in table.rows:
    for cell in row.cells:
        clear_all_borders(cell)
apply_three_line_borders(table, header_row=0, last_row=n_rows - 1, n_cols=n_cols)

doc.save(OUTPUT)
print(f"Saved: {OUTPUT}")
