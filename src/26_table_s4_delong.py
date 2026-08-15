"""
Table S4 (Appendix): Pairwise DeLong's test results by cohort, Holm-corrected
p-values - all 30 comparisons (6 per cohort x 5 cohorts), read directly from
the actual DeLong output workbooks (not retyped/recalled):
  outputs/delong_test_results.xlsx        (1-Year, 5-Year, Serial Biopsy)
  outputs/esrd/esrd_delong_results.xlsx   (ESRD 5-Year, ESRD 10-Year)

Three-line rule style (rule above header, below header, at foot; no vertical
rules, no grid) - matching the main-manuscript Table 1/2 convention
(src/25_tables_1_2.py), Times New Roman. Appendix tables were previously
built full-grid to match an older draft of the appendix; the appendix has
since been brought in line with the main-text style, so this table was
updated to match.

The single significant result across all 30 comparisons (Random Forest vs
LightGBM, 1-Year, Holm p=0.006) is bolded for visibility.

Saves: outputs/Table_S4_DeLong.docx
"""
import pandas as pd
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

BASE = "/Users/elizabetharmstrong/Library/CloudStorage/OneDrive-ImperialCollegeLondon/Lupus Project"
OUT = f"{BASE}/outputs"
OUTPUT = f"{OUT}/Table_S4_DeLong.docx"
FONT = "Times New Roman"

# --- Load real data (not retyped) ---
flare = pd.read_excel(f"{OUT}/delong_test_results.xlsx", sheet_name="All Cohorts")
esrd = pd.read_excel(f"{OUT}/esrd/esrd_delong_results.xlsx", sheet_name="All")
df = pd.concat([flare, esrd], ignore_index=True)

COHORT_ORDER = ["1-Year", "5-Year", "Serial Biopsy", "ESRD 5-Year", "ESRD 10-Year"]
df["Cohort"] = pd.Categorical(df["Cohort"], categories=COHORT_ORDER, ordered=True)
df = df.sort_values("Cohort", kind="stable").reset_index(drop=True)  # default quicksort isn't stable - scrambled within-cohort row order

print(df.to_string(index=False))
n_significant = (df["Significant (p<0.05)"] == "Yes").sum()
print(f"\n{len(df)} pairwise comparisons total, {n_significant} significant.")


# --- Border / font helpers (matching src/25_tables_1_2.py exactly) ---

def set_cell_borders(cell, top=None, bottom=None, left=None, right=None, insideH=None, insideV=None):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    specs = {"top": top, "bottom": bottom, "left": left, "right": right,
             "insideH": insideH, "insideV": insideV}
    for edge, spec in specs.items():
        el = tcBorders.find(qn(f'w:{edge}'))
        if el is None:
            el = OxmlElement(f'w:{edge}')
            tcBorders.append(el)
        if spec is None:
            el.set(qn('w:val'), 'nil')
        else:
            el.set(qn('w:val'), 'single')
            el.set(qn('w:sz'), str(spec))
            el.set(qn('w:color'), '000000')
            el.set(qn('w:space'), '0')


def clear_all_borders(cell):
    set_cell_borders(cell)


def apply_three_line_borders(table, header_row, last_row, n_cols):
    THICK, THIN = 18, 8
    # When header_row == 0, the top-rule and header-bottom-rule cells are the
    # SAME cell - set_cell_borders() rebuilds all 6 edges every call, so two
    # separate calls would let the second silently overwrite the first's top
    # edge back to nil. Combine into one call whenever they coincide (caught
    # via docx readback verification: row-0 top was rendering as 'nil').
    if header_row == 0:
        for c in range(n_cols):
            set_cell_borders(table.cell(0, c), top=THICK, bottom=THIN)
    else:
        for c in range(n_cols):
            set_cell_borders(table.cell(0, c), top=THICK)
        for c in range(n_cols):
            set_cell_borders(table.cell(header_row, c), bottom=THIN)
    for c in range(n_cols):
        set_cell_borders(table.cell(last_row, c), bottom=THICK)


def set_table_cell_margins(table, top_pt=3, bottom_pt=3, left_pt=5, right_pt=5):
    tbl = table._tbl
    tblPr = tbl.tblPr
    cellMar = OxmlElement('w:tblCellMar')
    for edge, pt_val in (('top', top_pt), ('bottom', bottom_pt), ('left', left_pt), ('right', right_pt)):
        node = OxmlElement(f'w:{edge}')
        node.set(qn('w:w'), str(int(pt_val * 20)))
        node.set(qn('w:type'), 'dxa')
        cellMar.append(node)
    tblPr.append(cellMar)


def set_run_font(run, size=10, bold=False, italic=False):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
        rFonts.set(qn(attr), FONT)


def set_cell(cell, text, size=10, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ""
    p = cell.paragraphs[0]
    for r in list(p.runs):  # cell.text = "" can leave a stray unformatted empty run
        r._element.getparent().remove(r._element)
    p.alignment = align
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic)


doc = Document()
section = doc.sections[0]
section.top_margin = section.bottom_margin = Cm(2.5)
section.left_margin = section.right_margin = Cm(2.5)

title = doc.add_paragraph()
run = title.add_run("Table S4. Pairwise DeLong's test results by cohort (Holm-corrected p-values).")
set_run_font(run, size=11, bold=True)
title.paragraph_format.space_after = Pt(8)

headers = ["Cohort", "Model A", "Model B", "AUROC A", "AUROC B", "p (raw)", "p (Holm)", "Significant"]
n_cols = len(headers)
n_rows = 1 + len(df)
table = doc.add_table(rows=n_rows, cols=n_cols)
table.style = "Table Grid"
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False
set_table_cell_margins(table, top_pt=3, bottom_pt=3, left_pt=5, right_pt=5)

col_widths = [2.6, 3.4, 3.4, 2.0, 2.0, 1.9, 1.9, 2.2]
for i, w in enumerate(col_widths):
    table.columns[i].width = Cm(w)
for row in table.rows:
    for cell, w in zip(row.cells, col_widths):
        cell.width = Cm(w)

for c, h in enumerate(headers):
    set_cell(table.cell(0, c), h, size=10, bold=True)

for r, row in enumerate(df.itertuples(index=False), start=1):
    is_sig = row._7 == "Yes"  # "Significant (p<0.05)" column
    set_cell(table.cell(r, 0), str(row.Cohort), align=WD_ALIGN_PARAGRAPH.LEFT, bold=is_sig)
    set_cell(table.cell(r, 1), row._1, align=WD_ALIGN_PARAGRAPH.LEFT, bold=is_sig)          # Model A
    set_cell(table.cell(r, 2), row._2, align=WD_ALIGN_PARAGRAPH.LEFT, bold=is_sig)          # Model B
    set_cell(table.cell(r, 3), f"{row._3:.3f}", bold=is_sig)                                 # AUROC A
    set_cell(table.cell(r, 4), f"{row._4:.3f}", bold=is_sig)                                 # AUROC B
    set_cell(table.cell(r, 5), f"{row._5:.3f}", bold=is_sig)                                 # p-value (DeLong)
    set_cell(table.cell(r, 6), f"{row._6:.3f}", bold=is_sig)                                 # p-value (Holm)
    set_cell(table.cell(r, 7), str(row._7), bold=is_sig)                                     # Significant

for row in table.rows:
    for cell in row.cells:
        clear_all_borders(cell)
apply_three_line_borders(table, header_row=0, last_row=n_rows - 1, n_cols=n_cols)

footnote = doc.add_paragraph()
run = footnote.add_run(
    f"Bonferroni-Holm correction applied across the 6 pairwise comparisons within each cohort "
    f"(30 comparisons total across 5 cohorts). One significant result: Random Forest significantly "
    f"outperformed LightGBM in the 1-Year cohort (p=0.001 raw, p=0.006 Holm-corrected), bolded above. "
    f"AUROC values are pooled out-of-fold (OOF) predictions from 5x10-fold repeated stratified CV "
    f"(5x5-fold for Serial Biopsy)."
)
set_run_font(run, size=9, italic=True)
footnote.paragraph_format.space_before = Pt(6)

doc.save(OUTPUT)
print(f"\nSaved: {OUTPUT}")
