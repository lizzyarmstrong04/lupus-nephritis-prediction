"""
Table 2 — Calibration / Accuracy (Brier score + calibration slope), formatted
to match the existing publication Table 1 (AUROC): Times New Roman, three-line
table style (top rule / header rule / bottom rule only, no vertical or internal
horizontal lines, no cell shading), Model rows x cohort columns grouped under
'Flare Cohorts' and 'Kidney Failure (ESRD) Cohorts'.

Each cell: Brier score (bold if lowest in that column among the four main
classifiers) on line 1, calibration slope (bold if closest to 1.0 in that
column; smaller, non-bold font otherwise) on line 2. TabPFN v3 is an added
row below the four main classifiers, entirely italic, not eligible for bold
highlighting (consistent with its treatment in Table 1).

NB: Random Forest / ESRD 10-Year calibration slope corrected to 1.152 (the
value verified against outputs/esrd/esrd_model_results.xlsx and the Methods
doc, Section 7.5) - the figure supplied in chat (0.963) was XGBoost's value
for that cell, not Random Forest's.

Saves: outputs/Table2_Calibration.docx
"""
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL

OUTPUT = "/Users/elizabetharmstrong/Library/CloudStorage/OneDrive-ImperialCollegeLondon/Lupus Project/outputs/Table2_Calibration.docx"
FONT = "Times New Roman"

COHORTS = ["1-Year Flare", "5-Year Flare", "ESRD 5-Year", "ESRD 10-Year"]
MODELS = ["Logistic Regression", "Random Forest", "XGBoost", "LightGBM"]

# [Model][Cohort] = (brier, cal_slope)
DATA = {
    "Logistic Regression": [(0.220, 0.719), (0.232, 0.669), (0.178, 0.873), (0.167, 0.874)],
    "Random Forest":       [(0.211, 0.849), (0.227, 0.827), (0.164, 1.013), (0.139, 1.152)],
    "XGBoost":              [(0.224, 0.835), (0.239, 0.705), (0.178, 1.288), (0.155, 0.963)],
    "LightGBM":             [(0.224, 0.686), (0.229, 0.878), (0.171, 1.053), (0.140, 0.714)],
}
TABPFN = [(0.166, 0.747), (0.230, 0.734), (0.100, 0.891), (0.122, 0.904)]

# --- Determine per-column bold winners ---
best_brier = {}   # col -> model with lowest Brier
best_cal   = {}    # col -> model with cal slope closest to 1.0
for col in range(4):
    briers = {m: DATA[m][col][0] for m in MODELS}
    cals   = {m: DATA[m][col][1] for m in MODELS}
    best_brier[col] = min(briers, key=briers.get)
    best_cal[col]   = min(cals, key=lambda m: abs(cals[m] - 1.0))

print("Bold Brier (lowest) per column:", {COHORTS[c]: best_brier[c] for c in range(4)})
print("Bold Cal slope (closest to 1.0) per column:", {COHORTS[c]: best_cal[c] for c in range(4)})

# --- Border helpers ---

def set_cell_borders(cell, top=None, bottom=None, left=None, right=None, insideH=None, insideV=None):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    specs = {"top": top, "bottom": bottom, "left": left, "right": right,
             "insideH": insideH, "insideV": insideV}
    for edge, spec in specs.items():
        el = tcBorders.find(qn(f'w:{edge}'))
        if el is None:
            el = OxmlElement(f'w:{edge}')
            tcBorders.append(el)
        if spec is None:
            el.set(qn('w:val'), 'nil')
        else:
            el.set(qn('w:val'), 'single')
            el.set(qn('w:sz'), str(spec))
            el.set(qn('w:color'), '000000')
            el.set(qn('w:space'), '0')


def clear_all_borders(cell):
    set_cell_borders(cell)  # all None -> all nil


def set_run_font(run, size=10.5, bold=False, italic=False):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
        rFonts.set(qn(attr), FONT)


def set_cell_lines(cell, lines, align=WD_ALIGN_PARAGRAPH.CENTER):
    """lines: list of (text, size, bold, italic) tuples, one per paragraph line."""
    cell.text = ""
    for i, (text, size, bold, italic) in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        run = p.add_run(text)
        set_run_font(run, size=size, bold=bold, italic=italic)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


# --- Build document ---

doc = Document()
section = doc.sections[0]
section.top_margin = section.bottom_margin = Cm(2.5)
section.left_margin = section.right_margin = Cm(2.5)

title = doc.add_paragraph()
run = title.add_run("Table 2. Calibration performance (Brier score, calibration slope) by model and cohort")
set_run_font(run, size=11, bold=True)
title.paragraph_format.space_after = Pt(8)

n_rows = 2 + len(MODELS) + 1  # 2 header rows + 4 model rows + 1 TabPFN row
n_cols = 1 + len(COHORTS)
table = doc.add_table(rows=n_rows, cols=n_cols)
table.style = "Table Grid"  # baseline; every border overridden explicitly below
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False

col_widths = [3.6] + [3.1] * 4
for row in table.rows:
    for cell, w in zip(row.cells, col_widths):
        cell.width = Cm(w)

# Baseline: strip every border on every cell first
for row in table.rows:
    for cell in row.cells:
        clear_all_borders(cell)

# Row 0/1: "Model" header, vertically merged across both header rows
model_hdr = table.cell(0, 0).merge(table.cell(1, 0))
set_cell_lines(model_hdr, [("Model", 11, True, False)], align=WD_ALIGN_PARAGRAPH.LEFT)

# Row 0: group headers, each merged across its 2 cohort columns
flare_hdr = table.cell(0, 1).merge(table.cell(0, 2))
set_cell_lines(flare_hdr, [("Flare Cohorts", 11, True, False)])
esrd_hdr = table.cell(0, 3).merge(table.cell(0, 4))
set_cell_lines(esrd_hdr, [("Kidney Failure (ESRD) Cohorts", 11, True, False)])

# Row 1: cohort names
for c, name in enumerate(COHORTS, start=1):
    set_cell_lines(table.cell(1, c), [(name, 11, True, False)])

# Rows 2-5: model data
for r, model in enumerate(MODELS, start=2):
    set_cell_lines(table.cell(r, 0), [(model, 11, False, False)], align=WD_ALIGN_PARAGRAPH.LEFT)
    for c, cohort in enumerate(COHORTS):
        brier, cal = DATA[model][c]
        brier_bold = (best_brier[c] == model)
        cal_bold = (best_cal[c] == model)
        set_cell_lines(table.cell(r, c + 1), [
            (f"{brier:.3f}", 11, brier_bold, False),
            (f"{cal:.3f}", 9, cal_bold, False),
        ])

# Last row: TabPFN v3, fully italic, no bold highlighting
tabpfn_row = 2 + len(MODELS)
set_cell_lines(table.cell(tabpfn_row, 0), [("TabPFN v3 (Prior Labs)", 11, False, True)], align=WD_ALIGN_PARAGRAPH.LEFT)
for c, cohort in enumerate(COHORTS):
    brier, cal = TABPFN[c]
    set_cell_lines(table.cell(tabpfn_row, c + 1), [
        (f"{brier:.3f}", 11, False, True),
        (f"{cal:.3f}", 9, False, True),
    ])

# --- Three-line borders: top rule, header separator, bottom rule ---
THICK = 18  # ~2.25pt
THIN = 8    # ~1pt

for c in range(n_cols):
    set_cell_borders(table.cell(0, c), top=THICK)         # table top rule
    set_cell_borders(table.cell(1, c), bottom=THIN)        # header/body separator
for c in range(n_cols):
    set_cell_borders(table.cell(tabpfn_row, c), bottom=THICK)  # table bottom rule

footnote = doc.add_paragraph()
run = footnote.add_run(
    "Bold Brier score = lowest (best) in that column among the four main classifiers. "
    "Bold calibration slope = closest to the ideal value of 1.0 in that column among the four main "
    "classifiers (a separate criterion from lowest Brier; the two need not point to the same model). "
    "TabPFN v3 (italic) is shown for reference and not included in either bold comparison."
)
set_run_font(run, size=9, italic=True)
footnote.paragraph_format.space_before = Pt(6)

doc.save(OUTPUT)
print(f"\nSaved: {OUTPUT}")
