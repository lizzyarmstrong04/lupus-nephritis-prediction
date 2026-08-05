from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

OUTPUT = "/Users/elizabetharmstrong/Library/CloudStorage/OneDrive-ImperialCollegeLondon/Lupus Project/outputs/Lupus_Project_Methods.docx"

doc = Document()

section = doc.sections[0]
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(2.8)
section.right_margin  = Cm(2.8)

styles = doc.styles

def set_style(style_name, font_name="Calibri", size=11, bold=False, color=None):
    s = styles[style_name]
    s.font.name  = font_name
    s.font.size  = Pt(size)
    s.font.bold  = bold
    if color:
        s.font.color.rgb = RGBColor(*color)

set_style("Normal",    size=11)
set_style("Heading 1", size=15, bold=True, color=(31, 73, 125))
set_style("Heading 2", size=13, bold=True, color=(31, 73, 125))
set_style("Heading 3", size=11, bold=True, color=(68, 114, 196))

def heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after  = Pt(4)
    return p

def body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    p.style = doc.styles["Normal"]
    return p

def bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.6 + level * 0.6)
    p.paragraph_format.space_after = Pt(3)
    return p

def add_table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(10)
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  "BDD7EE")
        tcPr.append(shd)
    for r_idx, row in enumerate(rows):
        tr = t.rows[r_idx + 1]
        for c_idx, val in enumerate(row):
            cell = tr.cells[c_idx]
            cell.text = str(val)
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            if r_idx % 2 == 1:
                tc   = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd  = OxmlElement("w:shd")
                shd.set(qn("w:val"),   "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"),  "F2F2F2")
                tcPr.append(shd)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return t

# TITLE PAGE

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Lupus Nephritis Prediction Pipeline")
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(31, 73, 125)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = subtitle.add_run(
    "Statistical Analysis & Machine Learning Methods\n"
    "1-Year Flare · 5-Year Flare · Serial Biopsy · ESRD Prediction · TabPFN Benchmark"
)
run2.font.size = Pt(13)
run2.font.color.rgb = RGBColor(68, 114, 196)

doc.add_paragraph()
date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_p.add_run(f"Generated: {datetime.date.today().strftime('%d %B %Y')}").font.size = Pt(11)

doc.add_page_break()

# 1. PROJECT OVERVIEW

heading(doc, "1. Project Overview", 1)
body(doc,
    "This document describes five parallel machine learning prediction pipelines developed "
    "from a single lupus nephritis (LN) biopsy database (1,070 biopsy episodes, 207 variables): "
    "(i) 1-year flare prediction, (ii) 5-year flare prediction, (iii) a serial biopsy "
    "sub-analysis incorporating longitudinal change between biopsies, (iv) 5-year end-stage "
    "renal disease (ESRD) prediction, and (v) 10-year ESRD prediction. All flare analyses share "
    "the same modelling framework — standardised preprocessing, Multiple Imputation by Chained "
    "Equations (MICE), a multi-step feature selection pipeline, four machine learning classifiers, "
    "repeated stratified cross-validation, Harrell optimism-corrected bootstrap, SHAP explainability, "
    "and pairwise DeLong's test for model comparison. The ESRD analyses follow the same modelling "
    "and evaluation framework — the same 6-step feature selection pipeline (leakage removal, "
    "dominant-binary, low-variance, high-correlation, VIF, LASSO with EPV-10 cap), MICE "
    "imputation (done as a separate pre-step before modelling), and Harrell bootstrap are applied. "
    "An additional TabPFN in-context learning benchmark was run on the flare datasets (Section 10)."
)
body(doc,
    "The primary outcome in flare analyses is binary (flare/no flare). The ESRD outcome is a "
    "composite binary endpoint: RRT (dialysis or transplant, codes 3/5) or creatinine doubling "
    "(code 2). Class imbalance was addressed in all four models. Feature selection was constrained "
    "by the Events Per Variable (EPV-10) rule appropriate to each cohort's event count."
)

add_table(doc,
    ["Analysis", "n", "Events", "Event rate", "Final predictors", "CV scheme"],
    [
        ["1-Year flare",          "430", "99",  "23.0%", "9",             "5×10-fold"],
        ["5-Year flare",          "356", "166", "46.6%", "10",            "5×10-fold"],
        ["Serial biopsy (5-yr)",   "70",  "34", "48.6%", "2",             "5×5-fold"],
        ["ESRD 5-year",           "796", "112", "14.1%", "5",  "5×10-fold"],
        ["ESRD 10-year",          "796", "175", "22.0%", "17", "5×10-fold"],
    ],
    col_widths=[4, 1.5, 2, 2.5, 3.5, 2.5]
)

# 2. DATA PREPARATION

heading(doc, "2. Data Preparation and Cohort Filtering", 1)

heading(doc, "2.1 Source Data", 2)
body(doc,
    "The source file (Data/Raw/data_lupus.xlsx) contains 1,070 biopsy episodes across 207 "
    "variables, including demographic, histopathological, serological, and treatment data. "
    "Each row represents a single biopsy episode. Some patients contributed more than one row "
    "(serial biopsies). No explicit patient identifier was present in the dataset."
)

heading(doc, "2.2 1-Year Flare Cohort  (src/1_year/01_data_prep_1yr.py)", 2)
body(doc,
    "The 1-year flare outcome column was identified by keyword search and coerced to numeric. "
    "Biopsy episodes were excluded under two criteria:"
)
bullet(doc, "Non-responders (NR, coded 1; technically partial response but NR pattern, coded 2): 177 excluded. Flare is only a meaningful outcome in patients who first achieved a response.")
bullet(doc, "Missing or unusable 1-year outcome (inadequate follow-up): 463 excluded.")
body(doc,
    "The resulting cohort comprised 430 biopsy episodes (99 flare events, 331 non-events; "
    "23.0% event rate). Saved as lupus_1yr_flare_dataset.xlsx."
)

heading(doc, "2.3 5-Year Flare Cohort  (src/5_year/01_data_prep_5yr.py)", 2)
body(doc,
    "The same exclusion logic was applied for the 5-year outcome. Non-responders were removed, "
    "and episodes without ≥5 years of follow-up (unless flare occurred before 5 years) were "
    "excluded. The 5-year cohort comprised 356 biopsy episodes (166 flare events, 190 non-events; "
    "46.6% event rate). Saved as lupus_5yr_flare_dataset.xlsx."
)

add_table(doc,
    ["", "1-Year Cohort", "5-Year Cohort"],
    [
        ["Total biopsy episodes",            "1,070",  "1,070"],
        ["Excluded: non-responders",           "177",    "177"],
        ["Excluded: inadequate follow-up",     "463",    "537"],
        ["Analysable cohort",                  "430",    "356"],
        ["Flare events (outcome = 1)",          "99",    "166"],
        ["Non-flare controls (outcome = 0)",   "331",    "190"],
        ["Event rate",                        "23.0%",  "46.6%"],
    ],
    col_widths=[8, 4, 4]
)

heading(doc, "2.4 Serial Biopsy Sub-Cohort  (src/5_year/08_serial_biopsy_5yr.py)", 2)
body(doc,
    "A serial biopsy sub-analysis was conducted within the 5-year cohort. Because no explicit "
    "patient identifier existed in the dataset, a proxy patient identifier was constructed as "
    "the concatenation of date of birth and gender (DOB + gender). This produced 634 unique "
    "patient groups with a maximum group size of 8, consistent with the known maximum of 9 "
    "biopsies per patient in this database."
)
body(doc,
    "Patients with at least two biopsies in the 5-year cohort (n=70) were selected. For each "
    "serial patient, the most recent biopsy episode was used as the index biopsy, and delta "
    "features were computed as the difference between the most recent and immediately preceding "
    "biopsy (most recent − previous). The following delta features were derived:"
)
for feat in ["% chronic gloms", "% active gloms", "% gloms with necrosis",
             "% sclerosed gloms", "LN class (coded)", "eGFR (CKD-EPI, no ethnicity)",
             "Proteinuria (log-transformed uPCR)"]:
    bullet(doc, f"Δ {feat}")
body(doc,
    "Time between biopsies (months) was also computed. Time-normalised delta features (Δ / "
    "time_between_bx_months, i.e. change per month) were added alongside the raw deltas to "
    "capture the rate of change independently of the interval length. Division by zero "
    "(biopsies on the same date) was handled by assigning NaN."
)
body(doc,
    "Prev exposure to cyclo was dropped before modelling: 69 of 70 patients had this value "
    "missing in the serial dataset, making it unusable. The 10 single-biopsy predictors from "
    "the 5-year feature set were carried forward from the most recent biopsy. "
    "The serial dataset comprised 70 patients with 34 flare events (48.6% event rate). "
    "Mean time between biopsies: 39.8 months."
)

heading(doc, "2.5 ESRD Cohort  (src/esrd/01_esrd_modelling.py)", 2)
body(doc,
    "The ESRD analyses use the full raw database without the non-responder or follow-up "
    "exclusions applied in the flare cohorts. All 1,070 biopsy episodes were considered; "
    "episodes were excluded only if the 5- or 10-year outcome column was missing (NaN) or "
    "marked 'X' (insufficient follow-up). This produced a larger analytic cohort (n=796) "
    "compared to the flare cohorts, because no treatment-response filter was applied."
)
body(doc,
    "The ESRD outcome is a composite binary endpoint: RRT/dialysis (code 3), death on RRT "
    "(code 5), or doubling of baseline creatinine (code 2). Outcome codes 4 (death alone, "
    "competing event), 6, 7, 8 (eGFR decline not meeting the composite), and X/NaN were "
    "excluded from the analytic dataset. Code 1 (stable eGFR>80) served as the non-event control."
)
add_table(doc,
    ["", "ESRD 5-Year", "ESRD 10-Year"],
    [
        ["Total biopsy episodes",        "1,070",   "1,070"],
        ["Excluded: NaN / X outcome",      "274",     "274"],
        ["Analysable cohort",              "796",     "796"],
        ["ESRD events (codes 2, 3, 5)",    "112",     "175"],
        ["Non-events (code 1)",            "684",     "621"],
        ["Event rate",                   "14.1%",   "22.0%"],
    ],
    col_widths=[8, 4, 4]
)
body(doc,
    "Note: the ESRD cohort includes repeat biopsies from the same patient (not restricted to "
    "index biopsies), consistent with the modelling script's data loading approach. This differs "
    "from the exploratory data preparation script (src/esrd/01_data_prep_esrd.py), which "
    "restricts to index biopsies (biopsy number = 1) for cohort characterisation."
)

# 3. MISSINGNESS ANALYSIS AND MICE IMPUTATION

heading(doc, "3. Missingness Analysis and MICE Imputation", 1)

heading(doc, "3.1 Approach", 2)
body(doc,
    "Missingness was assessed after leakage columns had been removed. Columns were categorised "
    "by percentage missing. Multiple Imputation by Chained Equations (MICE) was applied using "
    "scikit-learn's IterativeImputer (BayesianRidge estimator, 10 iterations, random_state=42). "
    "Only numeric columns with 0–50% missingness were imputed; columns >50% missing were excluded "
    "from imputation and subsequently excluded from modelling."
)

heading(doc, "3.2 1-Year Cohort  (src/1_year/04_mice_imputation_1yr.py)", 2)
add_table(doc,
    ["Missingness category", "Threshold", "Columns"],
    [
        ["No missing",           "0%",        "46"],
        ["Imputed (MICE)",       "0–50%",      "9"],
        ["Excluded from model",  ">50%",      "19"],
    ],
    col_widths=[7, 3.5, 3]
)
body(doc,
    "After MICE imputation, the 1-year imputed dataset (lupus_1yr_imputed.xlsx) contained "
    "zero missing values across all retained numeric columns."
)

heading(doc, "3.3 5-Year Cohort  (src/5_year/03_mice_imputation_5yr.py)", 2)
add_table(doc,
    ["Missingness category", "Threshold", "Columns"],
    [
        ["No missing",           "0%",        "50"],
        ["Imputed (MICE)",       "0–50%",     "10"],
        ["Excluded from model",  ">50%",      "18"],
    ],
    col_widths=[7, 3.5, 3]
)

heading(doc, "3.4 Serial Biopsy  (src/5_year/10_serial_feature_selection_5yr.py)", 2)
body(doc,
    "For the serial biopsy dataset (70 patients, 24 predictor columns after metadata removal), "
    "missingness arose only in delta features where one or both biopsy values were missing. "
    "All 12 columns with missing values fell below 5% (3 missing of 70), except delta_egfr "
    "and delta_log_proteinuria (1 missing each). All were imputed with MICE. No column exceeded "
    "50% missing. After imputation, zero missing values remained."
)

heading(doc, "3.5 ESRD Cohort  (src/esrd/00_esrd_feature_selection.py)", 2)
body(doc,
    "MICE imputation for the ESRD analyses is performed as a dedicated pre-step within the "
    "feature selection script (src/esrd/00_esrd_feature_selection.py), prior to and separate "
    "from the modelling script. This matches the flare pipeline design. The IterativeImputer "
    "(BayesianRidge, max_iter=10, random_state=42) was fitted on the full analytic cohort "
    "after leakage columns and columns with >50% missing values were removed. All candidate "
    "feature columns were coerced to numeric using pd.to_numeric(errors='coerce') before "
    "imputation, since several histopathological columns are stored as object type in the "
    "source Excel. After imputation, zero missing values remained."
)

# 4. FEATURE SELECTION

heading(doc, "4. Feature Selection", 1)

heading(doc, "4.1 Pipeline Overview", 2)
body(doc,
    "An identical multi-step feature selection pipeline was applied to the flare cohorts. "
    "Steps were applied in sequence; the output of each step was passed to the next:"
)
bullet(doc, "Step 1 — Leakage removal: any variable that encodes post-baseline information (outcomes, follow-up measurements, time-to-event, response dates) was explicitly excluded by name or keyword")
bullet(doc, "Step 2 — Date/text/ID exclusion: non-numeric columns")
bullet(doc, "Step 3a — Dominant binary removal: binary variables where one class >90% of data (near-zero discriminatory variance)")
bullet(doc, "Step 3b — Low-variance removal: numeric variables with variance <0.01")
bullet(doc, "Step 4 — High-correlation removal: pairwise Pearson r >0.80; the variable appearing first in the correlation matrix was removed iteratively")
bullet(doc, "Step 5 — VIF removal: Variance Inflation Factor >10 (after StandardScaler); iterative removal of worst variable")
bullet(doc, "Step 6 — LASSO regularisation: L1-penalised logistic regression (solver='saga') with decreasing C until ≤EPV_MAX features remained non-zero")

heading(doc, "4.2 Events Per Variable (EPV) Rule", 2)
body(doc,
    "The EPV-10 rule was applied to set the maximum number of predictors (EPV_MAX) for each cohort: "
    "EPV_MAX = floor(events / 10). This was used as the hard cap in the LASSO step."
)
add_table(doc,
    ["Cohort", "Events", "EPV_MAX", "Final predictors", "Actual EPV"],
    [
        ["1-Year flare",   "99",  "9",  "9",             "11.0"],
        ["5-Year flare",  "166", "15", "10",             "16.6"],
        ["Serial biopsy",  "34",  "4",  "2",             "17.0"],
        ["ESRD 5-Year",   "112", "11", "5",  "22.4"],
        ["ESRD 10-Year",  "175", "17", "17", "10.3"],
    ],
    col_widths=[4, 2.5, 3, 3.5, 3]
)
body(doc,
    "Note: the serial biopsy analysis ended with 2 predictors (EPV=17), not because the cap was "
    "raised, but because the LASSO path jumped from 12 features at C=0.5 directly to 2 features "
    "at C=0.2 with no intermediate value achieving 3–4 features. The two retained features are "
    "those with the strongest regularisation-robust association with the outcome in this small sample."
)

heading(doc, "4.3 1-Year Feature Selection  (src/1_year/05_feature_selection_1yr.py)", 2)

body(doc, "Step 1 — Leakage removal: 73 columns removed. Categories removed included:")
bullet(doc, "1-year creatinine, eGFR, proteinuria (measured at the outcome timepoint)")
bullet(doc, "Change-from-baseline columns (require knowing the 1-year value)")
bullet(doc, "Treatment response outcomes (CR, PR, NR, time to response, remission dates)")
bullet(doc, "Long-term renal and survival outcomes (RRT, death, CKD, creatinine doubling)")
bullet(doc, "Time-to-event outcomes (from diagnosis and from biopsy)")
bullet(doc, "Outcome-at-X-years columns (1, 5, 10 years)")
body(doc, "26 baseline numeric features remained after leakage and date/text exclusion.")

body(doc, "Step 3a — Dominant binary removal:")
add_table(doc,
    ["Variable removed", "Majority class (%)"],
    [["ANA (ever)", "93.7%"]],
    col_widths=[12, 4]
)

body(doc, "Step 3b — Low-variance removal:")
add_table(doc,
    ["Variable removed", "Variance"],
    [["Minimum of Creat/k and 1", "0.004"]],
    col_widths=[12, 4]
)

body(doc, "Step 4 — High-correlation removal (r > 0.80), 5 variables removed:")
add_table(doc,
    ["Variable removed", "Retained variable", "r"],
    [
        ["Age at biopsy",                        "Age (Now)",                            "0.933"],
        ["Age at diagnosis LN",                  "Age (Now)",                            "0.859"],
        ["Creat/0.9 if male or 0.7 if female",   "Creatinine at biopsy (mg/dl)",         "0.985"],
        ["Maximum of Creat/k and 1",             "Creatinine at biopsy (mg/dl)",         "0.984"],
        ["CKD-EPI with ethnicity",               "CKD-EPI without ethnicity",            "0.984"],
    ],
    col_widths=[6, 7, 2]
)

body(doc, "Step 5 — VIF removal: No variables exceeded VIF = 10. Maximum VIF = 3.41. No removals.")

body(doc, "Step 6 — LASSO (EPV_MAX = 9): C = 0.2 reduced 20 features to 10 non-zero. 10 variables zeroed:")
add_table(doc,
    ["Variable zeroed by LASSO at C=0.2"],
    [
        ["Creatinine at biopsy (mg/dl)"],
        ["CKD-EPI formula without ethnicity"],
        ["ANCA tested within 6 months of biopsy"],
        ["Age (Now)"],
        ["Biopsy number for patient"],
        ["Gender"],
        ["dsDNA / SM / APL ever positive"],
        ["Class change (yes/no)"],
        ["Number of biopsies for patient"],
        ["ANA (ever)"],
    ],
    col_widths=[16]
)

body(doc, "Post-selection corrections (manual, applied after automated pipeline):")
bullet(doc, "Age (Now) → replaced by Age at biopsy. Age Now reflected age at database compilation (~13.5 years post-biopsy on average) — not a valid baseline predictor.")
bullet(doc, "ANA (ever) → removed on clinical grounds (93.7% positive; near-zero discriminatory value even though not formally zeroed by LASSO at this C).")
bullet(doc, "ANCA tested within 6 months → removed. Originally 41.2% missing; MICE imputation of this binary column produced fractional values, making it unreliable.")
body(doc, "Candidate replacements were assessed by correlation with the existing feature set (r < 0.80) and VIF < 10:")
bullet(doc, "Biopsy number / ANA replacement: Proteinuria at biopsy (uPCR, log-transformed) — strongest correlation with outcome (r = 0.174), lowest VIF (1.19)")
bullet(doc, "ANCA replacement: C4 at biopsy — lowest max correlation with existing features (r = 0.290), VIF = 1.23")

body(doc, "Final 1-year predictors (9 variables, EPV = 11.0):")
add_table(doc,
    ["#", "Feature", "Category"],
    [
        ["1", "% chronic gloms (% of total)",             "Histopathology"],
        ["2", "% gloms with necrosis",                    "Histopathology"],
        ["3", "LN class (coded 1–10)",                    "Histopathology"],
        ["4", "% active gloms",                           "Histopathology"],
        ["5", "% gloms with crescents",                   "Histopathology"],
        ["6", "Age at biopsy",                            "Demographics"],
        ["7", "Ethnicity",                                "Demographics"],
        ["8", "Proteinuria at biopsy (uPCR, log-transformed)", "Renal function"],
        ["9", "C4 at biopsy",                             "Serology"],
    ],
    col_widths=[1, 8, 4]
)

heading(doc, "4.4 5-Year Feature Selection  (src/5_year/04_feature_selection_5yr.py)", 2)
body(doc,
    "The same pipeline was applied. The leakage list was expanded to additionally remove 1-year "
    "outcome variables (1-year proteinuria, 1-year creatinine/eGFR, 1-year steroid dose, "
    "'flare by 1 year'). 79 columns were removed at Step 1."
)

body(doc, "Step 3a: ANA (ever) removed (93.3% positive).")
body(doc, "Step 4 — High-correlation removal: Age at biopsy was removed (r=0.935 with Age Now). Further correlated pairs were removed iteratively.")
body(doc, "Step 5 — VIF removal: No variables exceeded VIF=10.")
body(doc,
    "Step 6 — LASSO (EPV_MAX = 15): C = 0.2 retained 13 non-zero features. "
    "Post-selection corrections:"
)
bullet(doc, "Removed Age (Now) → replaced by Age at biopsy (clinically appropriate baseline predictor)")
bullet(doc, "Removed Biopsy number for patient (proxy for disease duration)")
bullet(doc, "Removed Age at diagnosis SLE (r=0.826 with Age at biopsy — would fail correlation threshold)")
bullet(doc, "Removed dsDNA/SM/APL ever positive: 85.1% positive (just below 90% threshold), but clinical reasoning mirrors ANA removal — insufficient discriminatory gradient to justify inclusion")
body(doc, "Final 5-year predictors (10 variables, EPV = 16.6):")
add_table(doc,
    ["#", "Feature", "Category"],
    [
        ["1",  "% chronic gloms (% of total)",               "Histopathology"],
        ["2",  "% gloms with necrosis",                      "Histopathology"],
        ["3",  "LN class (coded 1–10)",                      "Histopathology"],
        ["4",  "% active gloms",                             "Histopathology"],
        ["5",  "% sclerosed gloms",                          "Histopathology"],
        ["6",  "Age at biopsy",                              "Demographics"],
        ["7",  "Ethnicity",                                  "Demographics"],
        ["8",  "Reason for biopsy (coded 1–4)",              "Clinical context"],
        ["9",  "CKD-EPI eGFR (without ethnicity)",           "Renal function"],
        ["10", "Prev exposure to cyclophosphamide",          "Treatment history"],
    ],
    col_widths=[1, 8, 4]
)

heading(doc, "4.5 Serial Biopsy Feature Selection  (src/5_year/10_serial_feature_selection_5yr.py)", 2)
body(doc,
    "The serial dataset entered the pipeline with 24 predictor columns (7 raw deltas, "
    "7 time-normalised deltas, time_between_bx_months, and 9 single-biopsy predictors). "
    "EPV_MAX = 4 (34 events)."
)
body(doc, "Step 3b: delta_ln_class_per_month removed (variance = 0.003).")
body(doc, "Step 4 — High-correlation removal: No pairs exceeded r = 0.80.")
body(doc, "Step 5 — VIF removal: % sclerosed gloms removed (VIF = 11.63). Max VIF after removal: 6.42.")
body(doc,
    "Step 6 — LASSO (EPV_MAX = 4): C search printed explicitly. The path showed a sharp "
    "discontinuity — 12 features at C=0.5, then 2 features at C=0.2 — with no intermediate "
    "C yielding 3 or 4 features. C=0.2 was selected as the first C achieving ≤4 features."
)
add_table(doc,
    ["C value", "Non-zero features"],
    [
        ["1.0",    "15"],
        ["0.5",    "12"],
        ["0.2",     "2  ← selected"],
        ["0.1",     "1"],
        ["≤0.05",   "0"],
    ],
    col_widths=[4, 5]
)
body(doc, "Final serial biopsy predictors (2 variables, EPV = 17.0):")
add_table(doc,
    ["#", "Feature", "LASSO coefficient"],
    [
        ["1", "% chronic gloms (% of total)",  "+0.291  (higher chronicity → higher flare risk)"],
        ["2", "Time between biopsies (months)", "−0.042  (longer interval → lower flare risk)"],
    ],
    col_widths=[1, 6, 8]
)

heading(doc, "4.6 ESRD Feature Selection  (src/esrd/00_esrd_feature_selection.py)", 2)
body(doc,
    "The same 6-step automated pipeline was applied to both ESRD cohorts. After leakage removal, "
    "numeric coercion, and MICE imputation, the pipeline proceeded identically to the flare analyses. "
    "Additional leakage keywords specific to ESRD were added (e.g. 'at last follow', 'post bx', "
    "'creatinine 1 year', '1year', 'immunostaining') to catch follow-up measurements present in the "
    "raw data. The manual correction 'Age (Now)' → 'Age at biopsy' was also applied (same as flare)."
)
body(doc, "ESRD 5-Year — final 5 features (EPV = 22.4, EPV_MAX = 11):")
add_table(doc,
    ["#", "Feature", "Category"],
    [
        ["1", "Creatinine at biopsy",              "Renal function"],
        ["2", "%IFTA",                             "Histopathology"],
        ["3", "CKD-EPI eGFR (without ethnicity)",  "Renal function"],
        ["4", "% chronic gloms (% of total)",      "Histopathology"],
        ["5", "Subepithelial deposit category",    "Histopathology (EM)"],
    ],
    col_widths=[1, 8, 4]
)
body(doc,
    "LASSO selected 5 features — well within the EPV_MAX=11 cap (EPV=22.4). Features reflect "
    "established markers of chronic renal damage and immune complex deposition."
)
body(doc, "ESRD 10-Year — final 17 features (EPV = 10.3, EPV_MAX = 17):")
add_table(doc,
    ["#", "Feature", "Category"],
    [
        ["1",  "Age at biopsy",                        "Demographics"],
        ["2",  "%IFTA",                                "Histopathology"],
        ["3",  "Creatinine at biopsy",                 "Renal function"],
        ["4",  "% chronic gloms (% of total)",         "Histopathology"],
        ["5",  "Crescents (Yes/No)",                   "Histopathology"],
        ["6",  "CKD-EPI eGFR (without ethnicity)",     "Renal function"],
        ["7",  "TMA (Yes/No)",                         "Histopathology"],
        ["8",  "LN class (coded 1–10)",                "Histopathology"],
        ["9",  "C3 at biopsy",                         "Serology"],
        ["10", "C4 low (binary)",                      "Serology"],
        ["11", "Prior cyclophosphamide exposure",      "Treatment history"],
        ["12", "Gender",                               "Demographics"],
        ["13", "Subepithelial deposit category",       "Histopathology (EM)"],
        ["14", "Cap wall IgM",                         "Histopathology (IF)"],
        ["15", "No. globally sclerosed gloms",         "Histopathology"],
        ["16", "Biopsy number for patient",            "Clinical context"],
        ["17", "No. gloms with crescents",             "Histopathology"],
    ],
    col_widths=[1, 8, 4]
)
body(doc,
    "LASSO selected exactly 17 features — at the EPV_MAX cap (EPV=10.3). The larger feature "
    "set at 10 years reflects the greater number of events (175) and broader range of predictors "
    "that retain discriminatory value over a longer horizon."
)

# 5. CLASS IMBALANCE HANDLING

heading(doc, "5. Class Imbalance Handling", 1)
body(doc,
    "Class imbalance was addressed in all four classifiers across all analyses. "
    "The approach differed by algorithm:"
)
bullet(doc, "Logistic Regression, Random Forest, LightGBM: class_weight='balanced' — sklearn automatically scales class weights inversely proportional to class frequencies")
bullet(doc, "XGBoost: scale_pos_weight = (number of negative cases) / (number of positive cases) — equivalent mechanism for gradient boosting")
add_table(doc,
    ["Cohort", "Event rate", "LR / RF / LGBM", "XGBoost scale_pos_weight"],
    [
        ["1-Year flare",   "23.0%", "class_weight='balanced'", "331 / 99 = 3.34"],
        ["5-Year flare",   "46.6%", "class_weight='balanced'", "190 / 166 = 1.14"],
        ["Serial biopsy",  "48.6%", "class_weight='balanced'", "36 / 34 = 1.06"],
        ["ESRD 5-Year",    "14.1%", "class_weight='balanced'", "684 / 112 = 6.11"],
        ["ESRD 10-Year",   "22.0%", "class_weight='balanced'", "621 / 175 = 3.55"],
    ],
    col_widths=[3.5, 3, 5.5, 5]
)
body(doc,
    "The serial biopsy cohort was nearly balanced (48.6% event rate), so weighting had minimal "
    "effect. The 1-year cohort and ESRD 5-year cohort had the most pronounced imbalances, "
    "where class weighting was most important."
)

# 6. MODEL DEVELOPMENT

heading(doc, "6. Model Development and Evaluation", 1)

heading(doc, "6.1 Algorithms", 2)
body(doc,
    "Four classifiers were developed for each cohort, all wrapped in a StandardScaler pipeline "
    "to ensure features were on the same scale before fitting:"
)
bullet(doc, "Logistic Regression (sklearn) — linear, interpretable baseline; C=1.0 (no regularisation beyond class weighting)")
bullet(doc, "Random Forest (sklearn) — ensemble of decision trees with bagging")
bullet(doc, "XGBoost — gradient-boosted trees with L1/L2 regularisation")
bullet(doc, "LightGBM — leaf-wise gradient boosting with histogram-based splits")

heading(doc, "6.2 Hyperparameter Tuning", 2)
body(doc,
    "Logistic Regression was not tuned (C=1.0 fixed). Random Forest, XGBoost, and LightGBM "
    "were tuned using RandomizedSearchCV (n_iter=40 for 1yr/5yr/ESRD; n_iter=20 for serial). "
    "Scoring metric: AUROC. Inner CV: stratified k-fold "
    "(5-fold for 1yr/5yr/ESRD; 3-fold for serial due to small n)."
)
body(doc, "XGBoost was constrained to prevent overfitting observed in preliminary runs (apparent bootstrap AUROC = 1.000 on unconstrained grid):")
bullet(doc, "max_depth: 2–3 only")
bullet(doc, "min_child_weight: ≥5 (1yr/5yr/ESRD) or ≥3 (serial)")
bullet(doc, "subsample and colsample_bytree: ≤0.7")
bullet(doc, "reg_alpha (L1) and reg_lambda (L2): explicit non-zero ranges")

body(doc, "Best hyperparameters — 1-Year cohort:")
add_table(doc,
    ["Model", "Parameter", "Value"],
    [
        ["Random Forest", "n_estimators / max_depth / min_samples_leaf / max_features", "300 / 3 / 10 / sqrt"],
        ["XGBoost",       "n_estimators / max_depth / lr / subsample / colsample / reg_α / reg_λ / min_child_w", "100 / 3 / 0.01 / 0.5 / 0.6 / 1.0 / 1.0 / 10"],
        ["LightGBM",      "n_estimators / max_depth / lr / subsample / num_leaves / min_child_samples", "200 / 2 / 0.01 / 0.8 / 15 / 10"],
    ],
    col_widths=[4, 7.5, 5]
)

body(doc, "Best hyperparameters — 5-Year cohort:")
add_table(doc,
    ["Model", "Parameter", "Value"],
    [
        ["Random Forest", "n_estimators / max_depth / min_samples_leaf / max_features", "200 / 3 / 5 / 0.5"],
        ["XGBoost",       "n_estimators / max_depth / lr / subsample / colsample / reg_α / reg_λ / min_child_w", "100 / 2 / 0.01 / 0.5 / 0.5 / 2.0 / 10.0 / 5"],
        ["LightGBM",      "n_estimators / max_depth / lr / subsample / num_leaves / min_child_samples", "100 / 3 / 0.01 / 0.7 / 15 / 10"],
    ],
    col_widths=[4, 7.5, 5]
)

body(doc, "Best hyperparameters — Serial biopsy cohort:")
add_table(doc,
    ["Model", "Parameter", "Value"],
    [
        ["Random Forest", "n_estimators / max_depth / min_samples_leaf / max_features", "100 / 2 / 5 / sqrt"],
        ["XGBoost",       "n_estimators / max_depth / lr / subsample / colsample / reg_α / reg_λ / min_child_w", "100 / 2 / 0.01 / 0.5 / 0.7 / 2.0 / 5.0 / 3"],
        ["LightGBM",      "n_estimators / max_depth / lr / subsample / num_leaves / min_child_samples", "50 / 3 / 0.01 / 1.0 / 31 / 20"],
    ],
    col_widths=[4, 7.5, 5]
)

body(doc, "Best hyperparameters — ESRD cohorts (5-year and 10-year use the same search grid):")
add_table(doc,
    ["Model", "Parameter", "Search space"],
    [
        ["Random Forest", "n_estimators / max_depth / min_samples_leaf / max_features", "200/300/500 / 3/5/7/None / 5/10/20 / sqrt/0.5/0.7"],
        ["XGBoost",       "n_estimators / max_depth / lr / subsample / colsample / min_child_w / reg_α / reg_λ", "100/200 / 2/3 / 0.01/0.05 / 0.5/0.6 / 0.5/0.6 / 5/10/15 / 0.1/1/2 / 1/5/10"],
        ["LightGBM",      "n_estimators / max_depth / lr / subsample / num_leaves / min_child_samples", "100/200/300 / 2/3/4 / 0.01/0.05/0.1 / 0.7/0.8/1.0 / 15/31/63 / 10/20/30"],
    ],
    col_widths=[4, 7.5, 5]
)

heading(doc, "6.3 Cross-Validation Strategy", 2)
body(doc,
    "Repeated stratified k-fold cross-validation was used. The number of splits was chosen "
    "to ensure a minimum of approximately 14 test patients per fold (to obtain stable "
    "per-fold AUROC estimates):"
)
bullet(doc, "1-Year (n=430): 5 repeats × 10 folds = 50 evaluations; ~43 test patients per fold")
bullet(doc, "5-Year (n=356): 5 repeats × 10 folds = 50 evaluations; ~36 test patients per fold")
bullet(doc, "Serial biopsy (n=70): 5 repeats × 5 folds = 25 evaluations; ~14 test patients per fold (10-fold would give only 7)")
bullet(doc, "ESRD 5-year and 10-year (n=796): 5 repeats × 10 folds = 50 evaluations; ~80 test patients per fold")
body(doc,
    "Stratification ensured proportional outcome representation in each fold. "
    "Out-of-fold (OOF) predicted probabilities were averaged across repeats for calibration curves. "
    "Three metrics were computed per fold and averaged:"
)
bullet(doc, "AUROC — discrimination; 95% CI from fold-level percentiles (2.5th–97.5th)")
bullet(doc, "Brier Score — probabilistic accuracy (lower = better)")
bullet(doc, "Calibration Slope — logistic regression of observed outcome on log-odds of predicted probability; ideal = 1.0; <1 indicates over-dispersion (predictions too extreme)")

heading(doc, "6.4 Harrell Bootstrap Optimism Correction", 2)
body(doc,
    "1,000 bootstrap iterations were performed using Harrell's optimism-correction method to "
    "estimate bias-corrected (BC) AUROC and Brier score for the flare and serial biopsy analyses. "
    "For each bootstrap sample b:"
)
bullet(doc, "Fit model on bootstrap sample D_b (n sampled with replacement from original D)")
bullet(doc, "Evaluate on D_b → C_boot (bootstrap apparent performance)")
bullet(doc, "Evaluate same model on original D → C_orig (bootstrap test performance)")
bullet(doc, "Optimism_b = C_boot − C_orig")
body(doc,
    "Mean optimism across 1,000 iterations was subtracted from the apparent performance "
    "(model fitted on the full dataset) to yield the BC estimate. Optimism reflects the "
    "degree of overfitting: a large difference between apparent and BC AUROC indicates the "
    "model has exploited random structure in the training data. "
    "Harrell bootstrap was performed for all analyses including ESRD (src/esrd/01_esrd_modelling.py). "
    "The ESRD 10-year tree models showed the largest optimism corrections (RF: 0.064, LGBM: 0.067), "
    "reflecting the greater model complexity with 17 predictors and 175 events."
)

# 7. RESULTS

heading(doc, "7. Results", 1)

heading(doc, "7.1 1-Year Cohort Results", 2)
body(doc, "n=430, 99 events (23.0%), 9 predictors. Null model Brier score ≈ 0.177.")
add_table(doc,
    ["Model", "CV AUROC", "95% CI", "Brier", "Cal Slope", "BC AUROC"],
    [
        ["Logistic Regression", "0.708", "0.553–0.858", "0.220", "0.719", "0.709"],
        ["Random Forest",       "0.690", "0.549–0.834", "0.211", "0.849", "0.730"],
        ["XGBoost",             "0.674", "0.544–0.837", "0.224", "0.835", "0.704"],
        ["LightGBM",            "0.659", "0.532–0.818", "0.224", "0.686", "0.710"],
    ],
    col_widths=[4.5, 3, 3.5, 2.5, 3, 3]
)
body(doc,
    "Logistic Regression achieved the highest CV AUROC (0.708) and had the lowest optimism "
    "(BC AUROC ≈ CV AUROC = 0.709), confirming minimal overfitting. Random Forest had "
    "the best calibration slope (0.849, closest to 1.0) and the highest BC AUROC (0.730) "
    "due to larger apparent performance being partially corrected. Tree models (XGBoost, LightGBM) "
    "showed slightly higher apparent AUROCs on bootstrap samples, suggesting moderate overfitting. "
    "No model exceeded AUROC 0.75, consistent with the inherent difficulty of predicting flare "
    "from a single biopsy with a limited predictor set."
)

heading(doc, "7.2 5-Year Cohort Results", 2)
body(doc, "n=356, 166 events (46.6%), 10 predictors. Null model Brier score ≈ 0.249.")
add_table(doc,
    ["Model", "CV AUROC", "95% CI", "Brier", "Cal Slope", "BC AUROC"],
    [
        ["Logistic Regression", "0.673", "0.536–0.788", "0.232", "0.669", "0.670"],
        ["Random Forest",       "0.679", "0.505–0.809", "0.227", "0.827", "0.726"],
        ["XGBoost",             "0.673", "0.479–0.811", "0.239", "0.705", "0.685"],
        ["LightGBM",            "0.678", "0.518–0.792", "0.229", "0.878", "0.737"],
    ],
    col_widths=[4.5, 3, 3.5, 2.5, 3, 3]
)
body(doc,
    "All four models achieved nearly identical CV AUROCs (0.673–0.679). LightGBM had the "
    "best calibration slope (0.878) and BC AUROC (0.737). The 5-year cohort has a higher "
    "event rate than the 1-year cohort but slightly lower discrimination — reflecting that "
    "5-year outcomes are harder to predict from biopsy alone, and that the additional events "
    "come from a broader range of trajectories. Confidence intervals are wider than 1-year "
    "due to the smaller cohort (n=356 vs n=430)."
)

heading(doc, "7.3 Serial Biopsy Results", 2)
body(doc,
    "n=70, 34 events (48.6%), 2 predictors. 5×5-fold CV. Results should be treated as "
    "exploratory given the very small sample size."
)
add_table(doc,
    ["Model", "CV AUROC", "95% CI", "Brier", "Cal Slope", "BC AUROC"],
    [
        ["Logistic Regression", "0.676", "0.434–0.873", "0.235", "0.570", "0.660"],
        ["Random Forest",       "0.588", "0.318–0.824", "0.251", "0.320", "0.656"],
        ["XGBoost",             "0.648", "0.453–0.818", "0.249", "0.053", "0.617"],
        ["LightGBM",            "0.631", "0.443–0.855", "0.240", "0.346", "0.594"],
    ],
    col_widths=[4.5, 3, 3.5, 2.5, 3, 3]
)
body(doc,
    "Logistic Regression was the best-performing model (CV AUROC 0.676). With only 2 predictors, "
    "LR is the natural choice — complex tree models offer no advantage. XGBoost calibration "
    "slope of 0.053 indicates near-constant predictions (severe miscalibration), making it "
    "unsuitable at this sample size. The very wide confidence intervals (e.g. 0.434–0.873 for LR) "
    "reflect the small n. These results confirm that serial biopsy data adds marginal discriminatory "
    "signal beyond the cross-sectional features, but the sample is insufficient for definitive conclusions."
)

heading(doc, "7.4 ESRD 5-Year Results  (src/esrd/01_esrd_modelling.py)", 2)
body(doc,
    "n=796, 112 events (14.1%), 5 predictors. 5×10-fold CV + Harrell bootstrap. "
    "ESRD prediction substantially outperforms flare prediction, reflecting that ESRD is driven "
    "by measurable chronic markers with stronger prognostic signal."
)
add_table(doc,
    ["Model", "CV AUROC", "95% CI", "Brier", "Cal Slope", "BC AUROC"],
    [
        ["Logistic Regression", "0.797", "0.669–0.904", "0.178", "0.873", "0.797"],
        ["Random Forest",       "0.787", "0.623–0.898", "0.164", "1.013", "0.803"],
        ["XGBoost",             "0.792", "0.631–0.901", "0.178", "1.288", "0.800"],
        ["LightGBM",            "0.797", "0.653–0.923", "0.171", "1.053", "0.808"],
    ],
    col_widths=[4.5, 3, 3.5, 2.5, 3, 3]
)
body(doc,
    "All four models achieved very similar CV AUROCs (0.787–0.797) with small Harrell optimism "
    "(0.007–0.047), indicating minimal overfitting. LightGBM had the highest BC AUROC (0.808). "
    "Random Forest achieved a near-perfect calibration slope (1.013). The compact 5-feature set "
    "(creatinine, %IFTA, eGFR, % chronic gloms, subepithelial deposits) captures the key "
    "histopathological and renal function markers driving ESRD risk, and no model shows a "
    "significant performance advantage over LR (DeLong all non-significant after Holm correction)."
)

heading(doc, "7.5 ESRD 10-Year Results  (src/esrd/01_esrd_modelling.py)", 2)
body(doc,
    "n=796, 175 events (22.0%), 17 predictors. 5×10-fold CV + Harrell bootstrap."
)
add_table(doc,
    ["Model", "CV AUROC", "95% CI", "Brier", "Cal Slope", "BC AUROC"],
    [
        ["Logistic Regression", "0.811", "0.656–0.903", "0.167", "0.874", "0.817"],
        ["Random Forest",       "0.817", "0.710–0.931", "0.139", "1.152", "0.897"],
        ["XGBoost",             "0.821", "0.696–0.933", "0.155", "0.963", "0.849"],
        ["LightGBM",            "0.809", "0.701–0.926", "0.140", "0.714", "0.926"],
    ],
    col_widths=[4.5, 3, 3.5, 2.5, 3, 3]
)
body(doc,
    "All four models performed well (CV AUROC 0.809–0.821). The larger optimism corrections at "
    "10 years (RF: 0.064, LGBM: 0.067) reflect the greater model complexity with 17 predictors. "
    "LightGBM achieved the highest BC AUROC (0.926) despite slightly lower CV AUROC, and its "
    "calibration slope (0.714) suggests mild under-prediction at high risk. The 10-year horizon "
    "shows marginally better discrimination than 5-year, consistent with more extreme outcomes "
    "(dialysis/death on RRT) being easier to discriminate from stable eGFR. No significant "
    "pairwise differences after Holm correction (see Section 8.3)."
)

# 8. STATISTICAL COMPARISON — DELONG'S TEST

heading(doc, "8. Statistical Comparison — DeLong's Test", 1)

heading(doc, "8.1 Method", 2)
body(doc,
    "DeLong's non-parametric test (DeLong et al., 1988, Biometrics 44:837–845) was used to "
    "compare AUROCs between all pairs of models within each cohort. The fast DeLong algorithm "
    "was implemented in Python using midrank statistics. Predicted probabilities were taken from "
    "out-of-fold CV predictions (averaged across repeats), ensuring paired comparisons on the "
    "same observations. Bonferroni-Holm correction was applied across the 6 pairwise tests "
    "within each cohort."
)

heading(doc, "8.2 Flare Cohort Results  (src/13_delong_test.py)", 2)
body(doc, "1-Year cohort pairwise comparisons (OOF AUROC: LR=0.704, RF=0.685, XGB=0.669, LGBM=0.658):")
add_table(doc,
    ["Model A", "Model B", "AUROC A", "AUROC B", "p (raw)", "p (Holm)", "Significant"],
    [
        ["Logistic Regression", "Random Forest",  "0.704", "0.685", "0.225", "0.450", "No"],
        ["Logistic Regression", "XGBoost",        "0.704", "0.669", "0.054", "0.214", "No"],
        ["Logistic Regression", "LightGBM",       "0.704", "0.658", "0.014", "0.070", "No"],
        ["Random Forest",       "XGBoost",        "0.685", "0.669", "0.109", "0.326", "No"],
        ["Random Forest",       "LightGBM",       "0.685", "0.658", "0.001", "0.006", "Yes"],
        ["XGBoost",             "LightGBM",       "0.669", "0.658", "0.438", "0.450", "No"],
    ],
    col_widths=[4, 4, 2.5, 2.5, 2.5, 2.5, 2.5]
)
body(doc,
    "One significant difference was identified: Random Forest outperformed LightGBM "
    "(p=0.001 raw, p=0.006 Holm-corrected). All other comparisons were non-significant. "
    "No model was significantly better than Logistic Regression after correction."
)

body(doc, "5-Year cohort: No significant pairwise differences after Bonferroni-Holm correction (all adjusted p ≥ 1.0). OOF AUROCs: LR=0.666, RF=0.677, XGB=0.666, LGBM=0.681.")
body(doc, "Serial biopsy cohort: No significant pairwise differences after correction (all adjusted p ≥ 0.26). Expected given n=70.")
body(doc,
    "Overall interpretation for flare: Across all three flare cohorts, there is no statistically "
    "significant difference in discrimination between the four classifiers, with the single "
    "exception of RF vs LightGBM in the 1-year cohort. Logistic Regression is at least as "
    "discriminatory as the tree-based approaches in all three settings, supporting its use as "
    "the primary model given its superior interpretability."
)

heading(doc, "8.3 ESRD Cohort Results  (src/esrd/03_esrd_delong.py)", 2)
body(doc, "ESRD 5-Year pairwise comparisons (OOF AUROC: LR=0.796, RF=0.780, XGB=0.786, LGBM=0.788):")
add_table(doc,
    ["Model A", "Model B", "AUROC A", "AUROC B", "p (raw)", "p (Holm)", "Significant"],
    [
        ["Logistic Regression", "Random Forest",  "0.796", "0.780", "0.134", "0.670", "No"],
        ["Logistic Regression", "XGBoost",        "0.796", "0.786", "0.284", "0.853", "No"],
        ["Logistic Regression", "LightGBM",       "0.796", "0.788", "0.576", "1.000", "No"],
        ["Random Forest",       "XGBoost",        "0.780", "0.786", "0.040", "0.239", "No"],
        ["Random Forest",       "LightGBM",       "0.780", "0.788", "0.213", "0.850", "No"],
        ["XGBoost",             "LightGBM",       "0.786", "0.788", "0.763", "1.000", "No"],
    ],
    col_widths=[4, 4, 2.5, 2.5, 2.5, 2.5, 2.5]
)
body(doc,
    "No significant pairwise differences at 5 years after Holm correction (all adjusted p ≥ 0.239). "
    "Models are essentially equivalent with only 5 features and 112 events. LR achieves the "
    "highest OOF AUROC (0.796), consistent with the small feature set favouring the linear model."
)

body(doc, "ESRD 10-Year pairwise comparisons (OOF AUROC: LR=0.810, RF=0.816, XGB=0.820, LGBM=0.816):")
add_table(doc,
    ["Model A", "Model B", "AUROC A", "AUROC B", "p (raw)", "p (Holm)", "Significant"],
    [
        ["Logistic Regression", "Random Forest",  "0.810", "0.816", "0.607", "1.000", "No"],
        ["Logistic Regression", "XGBoost",        "0.810", "0.820", "0.356", "1.000", "No"],
        ["Logistic Regression", "LightGBM",       "0.810", "0.816", "0.661", "1.000", "No"],
        ["Random Forest",       "XGBoost",        "0.816", "0.820", "0.522", "1.000", "No"],
        ["Random Forest",       "LightGBM",       "0.816", "0.816", "0.992", "1.000", "No"],
        ["XGBoost",             "LightGBM",       "0.820", "0.816", "0.736", "1.000", "No"],
    ],
    col_widths=[4, 4, 2.5, 2.5, 2.5, 2.5, 2.5]
)
body(doc,
    "No significant pairwise differences at 10 years (all adjusted p = 1.000). Models are "
    "indistinguishable in discrimination on OOF predictions. This contrasts with the previous "
    "pipeline (12 hardcoded features) where tree models significantly outperformed LR — suggesting "
    "the prior advantage was partly driven by the larger predictor set (tree models better exploit "
    "interaction effects with more features). With the formal EPV-10 constrained feature set, "
    "no classifier type holds a significant advantage in either ESRD horizon."
)

# 9. SHAP EXPLAINABILITY

heading(doc, "9. SHAP Explainability Analysis", 1)

heading(doc, "9.1 Method", 2)
body(doc,
    "SHapley Additive exPlanations (SHAP) values were computed for all four models on the full "
    "training dataset. SHAP decomposes each individual prediction into per-feature contributions, "
    "with positive values indicating increased predicted probability of the outcome. Two explainer types "
    "were used:"
)
bullet(doc, "Logistic Regression: shap.LinearExplainer (interventional feature perturbation)")
bullet(doc, "Random Forest, XGBoost, LightGBM: shap.TreeExplainer (tree path dependent). For binary classification, class-1 SHAP values were extracted from the 3D output array (shape: n_samples × n_features × n_classes).")
body(doc, "For each model, three plot types were produced:")
bullet(doc, "Beeswarm plot — distribution of SHAP values per feature, coloured by feature value (high=red, low=blue)")
bullet(doc, "Bar plot — mean absolute SHAP value per feature (global feature importance ranking)")
bullet(doc, "Waterfall plots — individual explanations for the highest-risk patient, lowest-risk patient, and decision-boundary patient (predicted probability closest to 0.5)")
body(doc, "A cross-model comparison bar chart was produced for each cohort.")

heading(doc, "9.2 1-Year SHAP Results  (src/1_year/07_shap_1yr.py)", 2)
body(doc, "Mean absolute SHAP values per feature per model (sorted by cross-model mean):")
add_table(doc,
    ["Feature", "Log Reg", "Rand Forest", "XGBoost", "LightGBM", "Mean"],
    [
        ["% chronic gloms",            "0.367", "0.078", "0.139", "0.262", "0.211"],
        ["Proteinuria at biopsy (log)","0.320", "0.056", "0.075", "0.181", "0.158"],
        ["% gloms with necrosis",      "0.247", "0.014", "0.001", "0.026", "0.072"],
        ["Age at biopsy",              "0.195", "0.024", "0.030", "0.031", "0.070"],
        ["LN class",                   "0.189", "0.017", "0.020", "0.012", "0.059"],
        ["Ethnicity",                  "0.191", "0.009", "0.015", "0.000", "0.054"],
        ["% active gloms",             "0.051", "0.012", "0.029", "0.046", "0.034"],
        ["C4 at biopsy",               "0.062", "0.010", "0.030", "0.010", "0.028"],
        ["% gloms with crescents",     "0.040", "0.008", "0.023", "0.000", "0.018"],
    ],
    col_widths=[4.5, 2.2, 2.8, 2.2, 2.8, 2]
)
body(doc,
    "% chronic gloms and proteinuria at biopsy were the two most important predictors across "
    "all models. Logistic Regression assigned substantially higher SHAP magnitudes than tree "
    "models, consistent with its linear decision boundary amplifying feature contributions. "
    "Tree models showed flatter importance distributions — RF in particular assigned very low "
    "SHAP to necrosis and ethnicity. LightGBM and XGBoost assigned near-zero importance to "
    "% crescents, suggesting this feature contributed mainly through the LR pathway."
)

heading(doc, "9.3 5-Year SHAP Results  (src/5_year/06_shap_5yr.py)", 2)
body(doc, "Mean absolute SHAP values per feature per model (sorted by cross-model mean):")
add_table(doc,
    ["Feature", "Log Reg", "Rand Forest", "XGBoost", "LightGBM", "Mean"],
    [
        ["% chronic gloms",            "0.322", "0.060", "0.072", "0.197", "0.163"],
        ["Age at biopsy",              "0.219", "0.048", "0.049", "0.166", "0.120"],
        ["LN class",                   "0.204", "0.022", "0.031", "0.070", "0.082"],
        ["% sclerosed gloms",          "0.092", "0.035", "0.066", "0.087", "0.070"],
        ["% active gloms",             "0.192", "0.014", "0.020", "0.052", "0.070"],
        ["% gloms with necrosis",      "0.155", "0.007", "0.001", "0.047", "0.052"],
        ["Prev exposure to cyclo",     "0.050", "0.026", "0.038", "0.064", "0.044"],
        ["Reason for biopsy",          "0.105", "0.008", "0.014", "0.000", "0.032"],
        ["CKD-EPI eGFR",               "0.047", "0.010", "0.007", "0.043", "0.027"],
        ["Ethnicity",                  "0.056", "0.002", "0.001", "0.010", "0.017"],
    ],
    col_widths=[4.5, 2.2, 2.8, 2.2, 2.8, 2]
)
body(doc,
    "% chronic gloms remained the most important predictor (mean |SHAP| = 0.163). Age at "
    "biopsy was the second most important in the 5-year model (0.120), more prominent than in "
    "the 1-year model, consistent with age being a stronger driver of long-term renal "
    "progression than of short-term inflammatory flare. LN class (0.082) and % active gloms "
    "(0.070) also contributed meaningfully. Ethnicity and eGFR were the least important features."
)

heading(doc, "9.4 ESRD SHAP Results  (src/esrd/02_esrd_shap.py)", 2)
body(doc, "ESRD 5-Year — Mean absolute SHAP values per feature per model (sorted by cross-model mean):")
add_table(doc,
    ["Feature", "Log Reg", "Rand Forest", "XGBoost", "LightGBM", "Mean"],
    [
        ["Creatinine at biopsy",      "0.334", "0.112", "0.340", "0.531", "0.329"],
        ["Subepithelial deposit cat.", "0.455", "0.040", "0.170", "0.336", "0.250"],
        ["eGFR (CKD-EPI)",            "0.247", "0.070", "0.226", "0.195", "0.184"],
        ["% chronic gloms",           "0.310", "0.037", "0.143", "0.205", "0.174"],
        ["%IFTA",                     "0.250", "0.042", "0.162", "0.108", "0.141"],
    ],
    col_widths=[4.5, 2.2, 2.8, 2.2, 2.8, 2]
)
body(doc,
    "Creatinine at biopsy was the dominant predictor for ESRD 5-year (mean |SHAP| = 0.329), "
    "with subepithelial deposit category (a marker of immune complex deposition pattern on "
    "electron microscopy) second (0.250). eGFR, % chronic gloms, and %IFTA were broadly similar "
    "in importance. With only 5 features, the model is compact and all features contribute "
    "meaningfully. This contrasts with the flare prediction SHAP results where histopathological "
    "activity markers dominated — in ESRD, chronic damage and renal function markers lead."
)

body(doc, "ESRD 10-Year — Mean absolute SHAP values per feature per model (top 10 shown, sorted by cross-model mean):")
add_table(doc,
    ["Feature", "Log Reg", "Rand Forest", "XGBoost", "LightGBM", "Mean"],
    [
        ["Creatinine at biopsy",   "0.192", "0.096", "0.511", "0.688", "0.372"],
        ["%IFTA",                  "0.363", "0.056", "0.313", "0.290", "0.255"],
        ["Age at biopsy",          "0.194", "0.026", "0.331", "0.464", "0.254"],
        ["% chronic gloms",        "0.391", "0.043", "0.255", "0.296", "0.246"],
        ["eGFR (CKD-EPI)",         "0.322", "0.053", "0.242", "0.278", "0.224"],
        ["C3 at biopsy",           "0.333", "0.019", "0.207", "0.286", "0.211"],
        ["C4 low",                 "0.380", "0.014", "0.227", "0.222", "0.211"],
        ["Crescents (Yes/No)",     "0.247", "0.017", "0.182", "0.196", "0.160"],
        ["Prior cyclophosphamide", "0.173", "0.012", "0.174", "0.278", "0.159"],
        ["Cap wall IgM",           "0.190", "0.013", "0.173", "0.201", "0.144"],
    ],
    col_widths=[4.5, 2.2, 2.8, 2.2, 2.8, 2]
)
body(doc,
    "Creatinine at biopsy was the top predictor at 10 years (mean |SHAP| = 0.372, driven "
    "strongly by XGBoost and LightGBM). %IFTA, age at biopsy, and % chronic gloms were "
    "broadly similar (0.246–0.255). Complement markers (C3, C4 low) featured prominently, "
    "reflecting the immunological component of longer-term ESRD risk beyond chronic damage "
    "alone. SHAP comparisons between 5-year and 10-year ESRD horizons are visualised in "
    "outputs/esrd/figures/esrd_shap_5yr_vs_10yr.png."
)

# 10. TABPFN BENCHMARK

heading(doc, "10. TabPFN Benchmark", 1)

heading(doc, "10.1 Method  (src/15_tabpfn_benchmark.py)", 2)
body(doc,
    "TabPFN (Hollmann et al., 2022) is a prior-data fitted network — a transformer pre-trained "
    "on millions of synthetic tabular datasets, performing in-context learning at inference time "
    "without any fitting to the target dataset during forward pass. Version 1 (tabpfn==0.1.11) "
    "was used; this version requires no account or licence, supports datasets with ≤1,000 training "
    "samples and ≤100 features, and runs on CPU. Configuration: N_ensemble_configurations=32, "
    "device='cpu', seed=42."
)
body(doc,
    "TabPFN was benchmarked on the 1-year and 5-year flare datasets using the same 5×10-fold "
    "repeated stratified CV protocol as the four main classifiers. Harrell optimism-corrected "
    "bootstrap was not performed for TabPFN: preliminary runs showed the bootstrap would require "
    "several hours of CPU time (each bootstrap iteration requires a new forward pass with the "
    "full synthetic-ensemble configuration), making it computationally infeasible. CV results "
    "only are reported. Note: the training set in each fold does not exceed 1,000 samples "
    "(1-year cohort: ~387 train, ~43 test; 5-year cohort: ~320 train, ~36 test), satisfying "
    "the TabPFN v1 constraint."
)

heading(doc, "10.2 Results", 2)
body(doc, "1-Year Flare — comparison with existing models (5×10-fold CV):")
add_table(doc,
    ["Model", "CV AUROC", "95% CI", "Brier", "Cal Slope"],
    [
        ["Logistic Regression", "0.708", "0.553–0.858", "0.220", "0.719"],
        ["Random Forest",       "0.690", "0.549–0.834", "0.211", "0.849"],
        ["XGBoost",             "0.674", "0.544–0.837", "0.224", "0.835"],
        ["LightGBM",            "0.659", "0.532–0.818", "0.224", "0.686"],
        ["TabPFN",              "0.705", "0.562–0.838", "0.165", "0.731"],
    ],
    col_widths=[4.5, 3, 3.5, 2.5, 3]
)
body(doc,
    "TabPFN matched Logistic Regression on AUROC (0.705 vs 0.708, within the confidence interval "
    "overlap). Its most notable advantage was Brier score: 0.165 vs 0.220 for LR and 0.211–0.224 "
    "for tree models. A lower Brier score indicates better-calibrated probability estimates "
    "overall. This suggests TabPFN's pre-trained prior produces more reliable probability "
    "outputs even without dataset-specific training."
)

body(doc, "5-Year Flare — comparison with existing models (5×10-fold CV):")
add_table(doc,
    ["Model", "CV AUROC", "95% CI", "Brier", "Cal Slope"],
    [
        ["Logistic Regression", "0.673", "0.536–0.788", "0.232", "0.669"],
        ["Random Forest",       "0.679", "0.505–0.809", "0.227", "0.827"],
        ["XGBoost",             "0.673", "0.479–0.811", "0.239", "0.705"],
        ["LightGBM",            "0.678", "0.518–0.792", "0.229", "0.878"],
        ["TabPFN",              "0.672", "0.491–0.816", "0.230", "0.691"],
    ],
    col_widths=[4.5, 3, 3.5, 2.5, 3]
)
body(doc,
    "At 5 years, TabPFN performed at parity with the existing models on all three metrics. "
    "AUROC (0.672) was within 0.007 of all comparators. Brier score (0.230) was in the middle "
    "of the range. The Brier advantage seen at 1 year did not replicate at 5 years, possibly "
    "because the higher event rate (46.6%) reduces the impact of probability calibration on "
    "the Brier score. Overall, TabPFN does not improve on the existing model suite and does "
    "not justify its additional computational overhead for this dataset."
)

# 11. FILE AND SCRIPT REFERENCE

heading(doc, "11. File and Script Reference", 1)

heading(doc, "11.1 Data Files", 2)
add_table(doc,
    ["File", "Location", "Description"],
    [
        ["data_lupus.xlsx",                    "Data/Raw/",       "Source data (1,070 episodes, 207 variables)"],
        ["lupus_1yr_flare_dataset.xlsx",       "Data/Processed/", "1-year cohort (430 patients)"],
        ["lupus_5yr_flare_dataset.xlsx",       "Data/Processed/", "5-year cohort (356 patients)"],
        ["lupus_1yr_imputed.xlsx",             "Data/Processed/", "MICE-imputed 1-year dataset"],
        ["lupus_5yr_imputed.xlsx",             "Data/Processed/", "MICE-imputed 5-year dataset"],
        ["lupus_1yr_selected_clean.xlsx",      "Data/Processed/", "Final 9 predictors, 1-year"],
        ["lupus_5yr_selected_clean.xlsx",      "Data/Processed/", "Final 10 predictors, 5-year"],
        ["lupus_5yr_serial_dataset.xlsx",      "Data/Processed/", "Serial biopsy dataset (70 patients, raw+delta features)"],
        ["lupus_5yr_serial_selected.xlsx",     "Data/Processed/", "Serial biopsy final 2 predictors"],
        ["lupus_esrd_dataset.xlsx",            "Data/Processed/", "ESRD cohort (index biopsies, from 01_data_prep_esrd.py)"],
        ["1yr_model_results.xlsx",             "outputs/",        "1-year CV, bootstrap, hyperparameters (3 sheets)"],
        ["1yr_best_params.json",               "outputs/",        "1-year best hyperparameters (JSON)"],
        ["5yr_model_results.xlsx",             "outputs/",        "5-year CV, bootstrap, hyperparameters (3 sheets)"],
        ["5yr_best_params.json",               "outputs/",        "5-year best hyperparameters (JSON)"],
        ["5yr_serial_model_results.xlsx",      "outputs/",        "Serial biopsy CV, bootstrap, hyperparameters"],
        ["shap_importance_table.xlsx",         "outputs/",        "Mean |SHAP| per feature per model — 1-year"],
        ["shap_importance_table_5yr.xlsx",     "outputs/",        "Mean |SHAP| per feature per model — 5-year"],
        ["delong_test_results.xlsx",           "outputs/",        "Pairwise DeLong's test, all flare cohorts (4 sheets)"],
        ["tabpfn_comparison.xlsx",             "outputs/",        "TabPFN vs existing models — 1yr and 5yr (2 sheets)"],
        ["esrd_model_results.xlsx",            "outputs/esrd/",   "ESRD CV results — 5yr and 10yr (2 sheets)"],
        ["esrd_shap_table_5yr.xlsx",           "outputs/esrd/",   "Mean |SHAP| per feature — ESRD 5-year"],
        ["esrd_shap_table_10yr.xlsx",          "outputs/esrd/",   "Mean |SHAP| per feature — ESRD 10-year"],
        ["esrd_delong_results.xlsx",           "outputs/esrd/",   "Pairwise DeLong's test — ESRD 5yr and 10yr"],
    ],
    col_widths=[5.5, 3.5, 7]
)

heading(doc, "11.2 Scripts", 2)
add_table(doc,
    ["Script", "Purpose"],
    [
        ["src/00_eda_raw.py",                          "EDA on full raw dataset"],
        ["src/1_year/01_data_prep_1yr.py",             "Filter raw data → 1-year cohort"],
        ["src/1_year/03_missingness_1yr.py",           "Missingness summary"],
        ["src/1_year/04_mice_imputation_1yr.py",       "MICE imputation (1-year)"],
        ["src/1_year/05_feature_selection_1yr.py",     "Feature selection pipeline (1-year)"],
        ["src/1_year/06_modelling_1yr.py",             "Tuning, 5×10-fold CV, Harrell bootstrap, plots (1-year)"],
        ["src/1_year/07_shap_1yr.py",                 "SHAP analysis — 1-year"],
        ["src/5_year/01_data_prep_5yr.py",             "Filter raw data → 5-year cohort"],
        ["src/5_year/03_mice_imputation_5yr.py",       "MICE imputation (5-year)"],
        ["src/5_year/04_feature_selection_5yr.py",     "Feature selection pipeline (5-year)"],
        ["src/5_year/05_modelling_5yr.py",             "Tuning, 5×10-fold CV, Harrell bootstrap, plots (5-year)"],
        ["src/5_year/06_shap_5yr.py",                 "SHAP analysis — 5-year"],
        ["src/5_year/08_serial_biopsy_5yr.py",        "Construct serial biopsy dataset (DOB+gender proxy ID, delta features)"],
        ["src/5_year/09_serial_prep_5yr.py",          "Drop cyclo, add time-normalised deltas"],
        ["src/5_year/10_serial_feature_selection_5yr.py", "Feature selection — serial biopsy (EPV_MAX=4)"],
        ["src/5_year/11_serial_modelling_5yr.py",     "Tuning, 5×5-fold CV, Harrell bootstrap — serial biopsy"],
        ["src/13_delong_test.py",                     "Pairwise DeLong's test, all flare cohorts"],
        ["src/14_shap_1yr_vs_5yr.py",                 "Comparison figure: 1-year vs 5-year SHAP feature importance"],
        ["src/15_tabpfn_benchmark.py",                "TabPFN benchmark (5×10-fold CV, 1yr and 5yr datasets)"],
        ["src/esrd/01_data_prep_esrd.py",             "ESRD cohort characterisation (index biopsies, exclusion breakdown)"],
        ["src/esrd/01_esrd_modelling.py",             "ESRD tuning, 5×10-fold CV, ROC/calibration plots (5yr and 10yr)"],
        ["src/esrd/02_esrd_shap.py",                  "SHAP analysis — ESRD 5-year and 10-year"],
        ["src/esrd/03_esrd_delong.py",                "Pairwise DeLong's test — ESRD"],
        ["src/esrd/04_esrd_shap_comparison.py",       "Comparison figure: ESRD 5yr vs 10yr SHAP"],
        ["src/esrd/05_esrd_vs_flare_shap.py",         "Comparison figure: ESRD vs flare SHAP importance"],
        ["src/esrd/06_esrd_delong_figure.py",         "Grouped bar chart with significance brackets — ESRD DeLong results"],
        ["src/generate_methods_doc.py",               "Generate this document (python-docx)"],
    ],
    col_widths=[7, 9]
)

heading(doc, "11.3 Software", 2)
add_table(doc,
    ["Package", "Version", "Use"],
    [
        ["Python",       "3.9",    "Language"],
        ["pandas",       "—",      "Data manipulation"],
        ["scikit-learn", "1.6.1",  "Preprocessing, MICE, LR, RF, CV, metrics"],
        ["xgboost",      "2.1.4",  "Gradient boosting (XGBoost)"],
        ["lightgbm",     "4.6.0",  "Gradient boosting (LightGBM)"],
        ["tabpfn",       "0.1.11", "TabPFN in-context learning benchmark (v1, no account)"],
        ["shap",         "0.49.1", "SHAP explainability"],
        ["statsmodels",  "—",      "VIF calculation"],
        ["scipy",        "—",      "DeLong's test (normal distribution)"],
        ["matplotlib",   "3.9.4",  "Plotting"],
        ["openpyxl",     "3.1.5",  "Excel I/O"],
        ["python-docx",  "—",      "Word document generation"],
    ],
    col_widths=[4, 3, 9]
)

doc.save(OUTPUT)
print(f"Saved: {OUTPUT}")
