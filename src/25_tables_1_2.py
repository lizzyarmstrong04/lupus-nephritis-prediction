"""
Single merged publication table (replaces the separate Table 1 + Table 2):
Times New Roman, three-line format (top rule, header rule, bottom rule only
- no internal vertical lines, no cell shading), landscape orientation.

ONE header block (Model / Flare Cohorts [1-Year, 5-Year, Serial Biopsy] /
Kidney Failure (ESRD) Cohorts [5-Year, 10-Year]), followed by three stacked
panels sharing it - no repeated header, no repeated "Brier"/"Cal" text in
every cell (the metric is named once in each panel's own label row instead):
  (a) Discrimination - Bias-Corrected AUROC (95% CI)   [bold = highest in column]
  (b) Brier Score                                      [bold = lowest in column]
  (c) Calibration Slope                                [bold = closest to 1.0 in column]

Panel (a) uses the Harrell bootstrap bias-corrected (BC) AUROC (with a
percentile 95% CI from the per-iteration bootstrap values), not CV AUROC -
BC AUROC is the paper's designated primary metric (Table S5's footnote,
TRIPOD guidance on optimism-corrected discrimination). See
src/33_bc_auroc_ci.py for the recomputation (the original per-cohort
scripts only saved the mean BC AUROC, no CI).

Each panel repeats the Model row labels (Logistic Regression, Random Forest,
XGBoost, LightGBM, then TabPFN v3 in italics - excluded from all bold
comparisons, not run through DeLong's test or Harrell bootstrap; TabPFN's
API-hosted nature makes a 1000-iteration bootstrap infeasible, so its AUROC
is CV AUROC, NOT the same bias-corrected convention as the other four rows
in panel (a) - flagged explicitly in the TabPFN footnote below). Serial
Biopsy's TabPFN v3 AUROC uses the fold-mean 'CV AUROC' convention (0.626) -
NOT the 0.588 pooled-OOF value used in the ROC/calibration figures
(src/16_roc_calibration_plots.py), a different aggregation of the same
underlying predictions.

One combined footnote block: TabPFN v3 exclusion note, and a serial-biopsy
caveat covering both the wide CIs (panel a) and degenerate calibration
(panels b/c) - self-contained (cites Riley et al. 2020 directly rather than
cross-referencing a specific section/file, so it stays correct wherever
this table is pasted).

Saves: outputs/Tables_1_2_Publication.docx
"""
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL

OUTPUT = "/Users/elizabetharmstrong/Library/CloudStorage/OneDrive-ImperialCollegeLondon/Lupus Project/outputs/Tables_1_2_Publication.docx"
FONT = "Times New Roman"

COHORTS = ["1-Year Flare", "5-Year Flare", "Serial Biopsy", "ESRD 5-Year", "ESRD 10-Year"]
FLARE_GROUP_SIZE = 3   # 1-Year, 5-Year, Serial Biopsy
ESRD_GROUP_SIZE  = 2   # ESRD 5-Year, ESRD 10-Year
MODELS = ["Logistic Regression", "Random Forest", "XGBoost", "LightGBM"]
N = len(COHORTS)
N_COLS = 1 + N

# --- Data: (auroc, ci_lower, ci_upper) / (brier, cal_slope) ---
# AUROC = Harrell bootstrap bias-corrected (BC) AUROC, not CV AUROC (switched
# from CV to BC - the paper's designated primary metric, per Table S5's
# footnote and TRIPOD reporting guidance). Values + 95% CI read from
# outputs/bc_auroc_ci.xlsx (src/33_bc_auroc_ci.py), which re-ran the Harrell
# bootstrap keeping per-iteration values for a percentile CI (the original
# per-cohort modelling scripts only kept the mean, so no CI existed before).
# 18/20 model x cohort combinations reproduced the already-published BC
# AUROC point estimate exactly; 5-Year Flare's Random Forest and XGBoost
# differed by <=0.005 (library-version drift since the original run, not a
# bug - see src/33_bc_auroc_ci.py's docstring) and the freshly-recomputed
# values are used here since they're paired with this table's own CI.
AUROC_DATA = {
    "Logistic Regression": [(0.709, 0.660, 0.760), (0.670, 0.620, 0.724), (0.660, 0.549, 0.770), (0.797, 0.754, 0.837), (0.817, 0.782, 0.853)],
    "Random Forest":       [(0.730, 0.692, 0.766), (0.725, 0.694, 0.754), (0.656, 0.573, 0.730), (0.803, 0.770, 0.837), (0.897, 0.878, 0.915)],
    "XGBoost":              [(0.704, 0.667, 0.741), (0.680, 0.641, 0.718), (0.617, 0.515, 0.712), (0.800, 0.765, 0.834), (0.849, 0.826, 0.872)],
    "LightGBM":             [(0.710, 0.670, 0.747), (0.737, 0.703, 0.769), (0.594, 0.490, 0.689), (0.808, 0.774, 0.841), (0.926, 0.904, 0.945)],
}
AUROC_TABPFN = [(0.684, 0.567, 0.834), (0.671, 0.508, 0.815), (0.626, 0.352, 0.841), (0.796, 0.621, 0.898), (0.817, 0.710, 0.926)]

CAL_DATA = {
    "Logistic Regression": [(0.220, 0.719), (0.232, 0.669), (0.235, 0.570), (0.178, 0.873), (0.167, 0.874)],
    "Random Forest":       [(0.211, 0.849), (0.227, 0.827), (0.251, 0.320), (0.164, 1.013), (0.139, 1.152)],
    "XGBoost":              [(0.224, 0.835), (0.239, 0.705), (0.249, 0.053), (0.178, 1.288), (0.155, 0.963)],
    "LightGBM":             [(0.224, 0.686), (0.229, 0.878), (0.240, 0.346), (0.171, 1.053), (0.140, 0.714)],
}
CAL_TABPFN = [(0.166, 0.747), (0.230, 0.734), (0.251, 0.190), (0.100, 0.891), (0.122, 0.904)]

# --- Bold winners ---
best_auroc = {}
for col in range(N):
    vals = {m: AUROC_DATA[m][col][0] for m in MODELS}
    top = max(vals.values())
    best_auroc[col] = {m for m in MODELS if vals[m] == top}

best_brier, best_cal = {}, {}
for col in range(N):
    briers = {m: CAL_DATA[m][col][0] for m in MODELS}
    cals   = {m: CAL_DATA[m][col][1] for m in MODELS}
    best_brier[col] = min(briers, key=briers.get)
    best_cal[col]   = min(cals, key=lambda m: abs(cals[m] - 1.0))

print("(a) bold AUROC (highest, ties included) per column:", {COHORTS[c]: sorted(best_auroc[c]) for c in range(N)})
print("(b) bold Brier (lowest) per column:", {COHORTS[c]: best_brier[c] for c in range(N)})
print("(c) bold Cal slope (closest to 1.0) per column:", {COHORTS[c]: best_cal[c] for c in range(N)})

# --- Border / font helpers ---

def set_cell_borders(cell, top=None, bottom=None, left=None, right=None, insideH=None, insideV=None):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    specs = {"top": top, "bottom": bottom, "left": left, "right": right,
             "insideH": insideH, "insideV": insideV}
    for edge, spec in specs.items():
        el = tcBorders.find(qn(f'w:{edge}'))
        if el is None:
            el = OxmlElement(f'w:{edge}')
            tcBorders.append(el)
        if spec is None:
            el.set(qn('w:val'), 'nil')
        else:
            el.set(qn('w:val'), 'single')
            el.set(qn('w:sz'), str(spec))
            el.set(qn('w:color'), '000000')
            el.set(qn('w:space'), '0')


def clear_all_borders(cell):
    set_cell_borders(cell)


def set_run_font(run, size=10.5, bold=False, italic=False):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
        rFonts.set(qn(attr), FONT)


CELL_SPACE_AFTER = Pt(3)


def set_cell_lines(cell, lines, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(0)):
    cell.text = ""
    for i, (text, size, bold, italic) in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_after = CELL_SPACE_AFTER
        p.paragraph_format.space_before = space_before if i == 0 else Pt(0)
        run = p.add_run(text)
        set_run_font(run, size=size, bold=bold, italic=italic)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_cell_run_line(cell, runs, align=WD_ALIGN_PARAGRAPH.CENTER):
    """runs: list of (text, size, bold, italic) tuples, all on ONE line."""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = CELL_SPACE_AFTER
    p.paragraph_format.space_before = Pt(0)
    for text, size, bold, italic in runs:
        run = p.add_run(text)
        set_run_font(run, size=size, bold=bold, italic=italic)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_table_cell_margins(table, top_pt=3, bottom_pt=3, left_pt=5, right_pt=5):
    tbl = table._tbl
    tblPr = tbl.tblPr
    cellMar = OxmlElement('w:tblCellMar')
    for edge, pt_val in (('top', top_pt), ('bottom', bottom_pt), ('left', left_pt), ('right', right_pt)):
        node = OxmlElement(f'w:{edge}')
        node.set(qn('w:w'), str(int(pt_val * 20)))
        node.set(qn('w:type'), 'dxa')
        cellMar.append(node)
    tblPr.append(cellMar)


def apply_three_line_borders(table, header_row, last_row, n_cols):
    THICK, THIN = 18, 8
    for c in range(n_cols):
        set_cell_borders(table.cell(0, c), top=THICK)
        set_cell_borders(table.cell(header_row, c), bottom=THIN)
    for c in range(n_cols):
        set_cell_borders(table.cell(last_row, c), bottom=THICK)


def prevent_page_breaks(table):
    """No native python-docx attribute for this (Row has no
    allow_break_across_pages property) - sets the underlying OOXML directly:
    <w:cantSplit/> on every row so a row can't split across a page break,
    and keep_with_next on every paragraph in every row except the last so
    Word doesn't insert a page break between rows either."""
    rows = table.rows
    for row in rows:
        trPr = row._tr.get_or_add_trPr()
        trPr.append(OxmlElement('w:cantSplit'))
    for row in rows[:-1]:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.keep_with_next = True


def add_footnote(doc, text, space_after=Pt(18)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=9, italic=True)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = space_after
    return p


TABPFN_FOOTNOTE = ("TabPFN v3 shown in italics for reference throughout; not included in formal "
                    "pairwise significance testing (DeLong's test) or bootstrap correction. Panel (a) "
                    "AUROC for the four main classifiers is Harrell bootstrap bias-corrected; TabPFN v3's "
                    "1000-iteration bootstrap refitting is infeasible via its hosted API, so its AUROC is "
                    "the cross-validation (CV) mean instead - not directly comparable to the other four "
                    "rows on a like-for-like basis.")
SERIAL_FOOTNOTE = ("Serial Biopsy (all three panels) should be treated as exploratory: n=70 (34 events) "
                    "produces wide confidence intervals in panel (a) and degenerate calibration in "
                    "panels (b)/(c) (e.g. XGBoost's slope of 0.053 reflects near-constant predictions, "
                    "not a data error) — this cohort is formally underpowered by standard minimum "
                    "sample-size criteria for prediction model development (Riley et al. 2020, BMJ "
                    "368:m441).")

# --- Build document ---

doc = Document()
section = doc.sections[0]
section.orientation = WD_ORIENT.LANDSCAPE
section.page_width, section.page_height = section.page_height, section.page_width
section.top_margin = section.bottom_margin = Cm(2.5)
section.left_margin = section.right_margin = Cm(2.5)

title = doc.add_paragraph()
run = title.add_run("Table 1. Model performance by cohort: discrimination, Brier score, and calibration slope")
set_run_font(run, size=11, bold=True)
title.paragraph_format.space_after = Pt(8)
title.paragraph_format.keep_with_next = True

PANEL_ROWS = 1 + len(MODELS) + 1  # label row + 4 models + TabPFN
n_rows = 2 + 3 * PANEL_ROWS       # 2 header rows + 3 panels
table = doc.add_table(rows=n_rows, cols=N_COLS)
table.style = "Table Grid"
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False
set_table_cell_margins(table, top_pt=3, bottom_pt=3, left_pt=5, right_pt=5)

col_widths = [3.8] + [2.62] * N
# Both are needed: table.columns[i].width sets the tblGrid (what determines
# layout for cells that don't override it); per-cell .width sets each
# cell's own tcW. Setting only one or the other silently fails to render
# the intended widths in Word (verified empirically - cell-only left every
# column at a uniform default width).
for i, w in enumerate(col_widths):
    table.columns[i].width = Cm(w)
for row in table.rows:
    for cell, w in zip(row.cells, col_widths):
        cell.width = Cm(w)

for row in table.rows:
    for cell in row.cells:
        clear_all_borders(cell)

# Shared header block (rows 0-1)
model_hdr = table.cell(0, 0).merge(table.cell(1, 0))
set_cell_lines(model_hdr, [("Model", 11, True, False)], align=WD_ALIGN_PARAGRAPH.LEFT)

flare_hdr = table.cell(0, 1)
for c in range(2, 1 + FLARE_GROUP_SIZE):
    flare_hdr = flare_hdr.merge(table.cell(0, c))
set_cell_lines(flare_hdr, [("Flare Cohorts", 11, True, False)])

esrd_start = 1 + FLARE_GROUP_SIZE
esrd_hdr = table.cell(0, esrd_start)
for c in range(esrd_start + 1, esrd_start + ESRD_GROUP_SIZE):
    esrd_hdr = esrd_hdr.merge(table.cell(0, c))
set_cell_lines(esrd_hdr, [("Kidney Failure (ESRD) Cohorts", 11, True, False)])

for c, name in enumerate(COHORTS, start=1):
    set_cell_lines(table.cell(1, c), [(name, 11, True, False)])


def fill_panel(start_row, label, cell_fn):
    """label row (merged across all columns) + 4 model rows + TabPFN row."""
    label_row = table.cell(start_row, 0).merge(table.cell(start_row, N_COLS - 1))
    set_cell_lines(label_row, [(label, 11, True, False)], align=WD_ALIGN_PARAGRAPH.LEFT, space_before=Pt(3))
    for r, model in enumerate(MODELS, start=start_row + 1):
        set_cell_lines(table.cell(r, 0), [(model, 11, False, False)], align=WD_ALIGN_PARAGRAPH.LEFT)
        for c in range(N):
            cell_fn(table.cell(r, c + 1), model, c, italic=False)
    tabpfn_row = start_row + 1 + len(MODELS)
    set_cell_lines(table.cell(tabpfn_row, 0), [("TabPFN v3", 10.5, False, True)], align=WD_ALIGN_PARAGRAPH.LEFT)
    for c in range(N):
        cell_fn(table.cell(tabpfn_row, c + 1), None, c, italic=True)
    return tabpfn_row


def auroc_cell(cell, model, c, italic):
    auroc, lo, hi = (AUROC_DATA[model][c] if model else AUROC_TABPFN[c])
    bold = (not italic) and (model in best_auroc[c])
    set_cell_run_line(cell, [
        (f"{auroc:.3f}", 11, bold, italic),
        (f" ({lo:.2f}–{hi:.2f})", 9, False, italic),
    ])


def brier_cell(cell, model, c, italic):
    brier, _ = (CAL_DATA[model][c] if model else CAL_TABPFN[c])
    bold = (not italic) and (best_brier[c] == model)
    set_cell_lines(cell, [(f"{brier:.3f}", 11, bold, italic)])


def cal_cell(cell, model, c, italic):
    _, cal = (CAL_DATA[model][c] if model else CAL_TABPFN[c])
    bold = (not italic) and (best_cal[c] == model)
    set_cell_lines(cell, [(f"{cal:.3f}", 11, bold, italic)])


row = 2
row = fill_panel(row, "(a) Discrimination — Bias-Corrected AUROC (95% CI)", auroc_cell) + 1
row = fill_panel(row, "(b) Brier Score", brier_cell) + 1
last_row = fill_panel(row, "(c) Calibration Slope", cal_cell)

apply_three_line_borders(table, header_row=1, last_row=last_row, n_cols=N_COLS)
prevent_page_breaks(table)

add_footnote(doc, TABPFN_FOOTNOTE)
add_footnote(doc, SERIAL_FOOTNOTE)

doc.save(OUTPUT)
print(f"\nSaved: {OUTPUT}")
