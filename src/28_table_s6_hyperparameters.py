"""
Table S6 (Appendix): Selected hyperparameter values by cohort and model,
read directly from the four saved best-params sources (not retyped):
  outputs/1yr_best_params.json
  outputs/5yr_best_params.json
  outputs/5yr_serial_model_results.xlsx  (sheet "Best Hyperparameters")
  outputs/esrd/esrd_best_params.json     (keys "5yr" / "10yr")

Logistic Regression is untuned throughout (C=1.0, class_weight='balanced')
so is not included - only Random Forest, XGBoost, LightGBM have a tuned
hyperparameter search.

Random Forest, XGBoost and LightGBM have different hyperparameter sets, so
one shared table would be mostly blank cells - instead three compact
sub-tables (one per model, cohort rows). Each sub-table uses the three-line
rule style (rule above header, below header, at foot; no vertical rules,
no grid) matching the main-manuscript Table 1/2 convention
(src/25_tables_1_2.py), Times New Roman.

Saves: outputs/Table_S6_Hyperparameters.docx
"""
import json
import pandas as pd
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT

BASE = "/Users/elizabetharmstrong/Library/CloudStorage/OneDrive-ImperialCollegeLondon/Lupus Project"
OUT = f"{BASE}/outputs"
OUTPUT = f"{OUT}/Table_S6_Hyperparameters.docx"
FONT = "Times New Roman"

COHORTS = ["1-Year Flare", "5-Year Flare", "Serial Biopsy", "ESRD 5-Year", "ESRD 10-Year"]

# --- Load real data (not retyped) ---
with open(f"{OUT}/1yr_best_params.json") as f:
    p_1yr = json.load(f)
with open(f"{OUT}/5yr_best_params.json") as f:
    p_5yr = json.load(f)
with open(f"{OUT}/esrd/esrd_best_params.json") as f:
    p_esrd = json.load(f)

serial_df = pd.read_excel(f"{OUT}/5yr_serial_model_results.xlsx", sheet_name="Best Hyperparameters")
p_serial = {}
for _, row in serial_df.iterrows():
    p_serial[row["Model"]] = {f"clf__{c}": row[c] for c in serial_df.columns
                               if c != "Model" and pd.notna(row[c])}

PARAMS_BY_COHORT = {
    "1-Year Flare": p_1yr,
    "5-Year Flare": p_5yr,
    "Serial Biopsy": p_serial,
    "ESRD 5-Year": p_esrd["5yr"],
    "ESRD 10-Year": p_esrd["10yr"],
}

INT_PARAMS = {"n_estimators", "max_depth", "min_samples_leaf", "min_child_weight",
              "num_leaves", "min_child_samples"}


def fmt(val):
    if pd.isna(val):
        return "-"
    if isinstance(val, str):
        return val
    if isinstance(val, float) and val == int(val):
        return str(int(val))
    return str(val)


def get(model, cohort, param):
    d = PARAMS_BY_COHORT[cohort].get(model, {})
    return fmt(d.get(f"clf__{param}"))


MODEL_SPECS = {
    "(a) Random Forest": ["n_estimators", "max_depth", "min_samples_leaf", "max_features"],
    "(b) XGBoost": ["n_estimators", "max_depth", "learning_rate", "subsample",
                    "colsample_bytree", "min_child_weight", "reg_alpha", "reg_lambda"],
    "(c) LightGBM": ["n_estimators", "max_depth", "learning_rate", "subsample",
                     "num_leaves", "min_child_samples"],
}
PARAM_LABELS = {
    "n_estimators": "n_estimators", "max_depth": "max_depth",
    "min_samples_leaf": "min_samples_\nleaf", "max_features": "max_features",
    "learning_rate": "learning_rate", "subsample": "subsample",
    "colsample_bytree": "colsample_\nbytree", "min_child_weight": "min_child_\nweight",
    "reg_alpha": "reg_alpha (L1)", "reg_lambda": "reg_lambda (L2)",
    "num_leaves": "num_leaves", "min_child_samples": "min_child_\nsamples",
}

for label, params in MODEL_SPECS.items():
    model_name = label.split(") ", 1)[1]
    print(f"--- {label} ---")
    for cohort in COHORTS:
        vals = [get(model_name, cohort, p) for p in params]
        print(f"  {cohort:<16} " + "  ".join(f"{p}={v}" for p, v in zip(params, vals)))


def set_cell_borders(cell, top=None, bottom=None, left=None, right=None, insideH=None, insideV=None):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    specs = {"top": top, "bottom": bottom, "left": left, "right": right,
             "insideH": insideH, "insideV": insideV}
    for edge, spec in specs.items():
        el = tcBorders.find(qn(f'w:{edge}'))
        if el is None:
            el = OxmlElement(f'w:{edge}')
            tcBorders.append(el)
        if spec is None:
            el.set(qn('w:val'), 'nil')
        else:
            el.set(qn('w:val'), 'single')
            el.set(qn('w:sz'), str(spec))
            el.set(qn('w:color'), '000000')
            el.set(qn('w:space'), '0')


def clear_all_borders(cell):
    set_cell_borders(cell)


def apply_three_line_borders(table, header_row, last_row, n_cols):
    THICK, THIN = 18, 8
    # When header_row == 0, the top-rule and header-bottom-rule cells are the
    # SAME cell - set_cell_borders() rebuilds all 6 edges every call, so two
    # separate calls would let the second silently overwrite the first's top
    # edge back to nil. Combine into one call whenever they coincide (caught
    # via docx readback verification: row-0 top was rendering as 'nil').
    if header_row == 0:
        for c in range(n_cols):
            set_cell_borders(table.cell(0, c), top=THICK, bottom=THIN)
    else:
        for c in range(n_cols):
            set_cell_borders(table.cell(0, c), top=THICK)
        for c in range(n_cols):
            set_cell_borders(table.cell(header_row, c), bottom=THIN)
    for c in range(n_cols):
        set_cell_borders(table.cell(last_row, c), bottom=THICK)


def set_table_cell_margins(table, top_pt=3, bottom_pt=3, left_pt=5, right_pt=5):
    tbl = table._tbl
    tblPr = tbl.tblPr
    cellMar = OxmlElement('w:tblCellMar')
    for edge, pt_val in (('top', top_pt), ('bottom', bottom_pt), ('left', left_pt), ('right', right_pt)):
        node = OxmlElement(f'w:{edge}')
        node.set(qn('w:w'), str(int(pt_val * 20)))
        node.set(qn('w:type'), 'dxa')
        cellMar.append(node)
    tblPr.append(cellMar)


def set_run_font(run, size=10, bold=False, italic=False):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
        rFonts.set(qn(attr), FONT)


def set_cell(cell, text, size=9.5, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ""
    p = cell.paragraphs[0]
    for r in list(p.runs):  # cell.text = "" can leave a stray unformatted empty run
        r._element.getparent().remove(r._element)
    p.alignment = align
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic)


doc = Document()
section = doc.sections[0]
section.orientation = WD_ORIENT.LANDSCAPE
section.page_width, section.page_height = section.page_height, section.page_width
section.top_margin = section.bottom_margin = Cm(2.5)
section.left_margin = section.right_margin = Cm(2.5)

title = doc.add_paragraph()
run = title.add_run("Table S6. Selected hyperparameter values by cohort and model.")
set_run_font(run, size=11, bold=True)
title.paragraph_format.space_after = Pt(4)

note = doc.add_paragraph()
run = note.add_run(
    "Logistic Regression is untuned throughout (C=1.0, class_weight='balanced') and is not shown. "
    "Random Forest, XGBoost, and LightGBM were tuned via RandomizedSearchCV (scoring=AUROC)."
)
set_run_font(run, size=9, italic=True)
note.paragraph_format.space_after = Pt(10)

for label, params in MODEL_SPECS.items():
    model_name = label.split(") ", 1)[1]

    sub = doc.add_paragraph()
    run = sub.add_run(label)
    set_run_font(run, size=10.5, bold=True)
    sub.paragraph_format.space_before = Pt(8)
    sub.paragraph_format.space_after = Pt(4)

    headers = ["Cohort"] + [PARAM_LABELS[p] for p in params]
    n_rows = 1 + len(COHORTS)
    table = doc.add_table(rows=n_rows, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_cell_margins(table, top_pt=3, bottom_pt=3, left_pt=5, right_pt=5)

    col_widths = [3.0] + [2.3] * len(params)
    for i, w in enumerate(col_widths):
        table.columns[i].width = Cm(w)
    for row in table.rows:
        for cell, w in zip(row.cells, col_widths):
            cell.width = Cm(w)

    for c, h in enumerate(headers):
        set_cell(table.cell(0, c), h, bold=True)

    for r, cohort in enumerate(COHORTS, start=1):
        set_cell(table.cell(r, 0), cohort, align=WD_ALIGN_PARAGRAPH.LEFT)
        for c, param in enumerate(params, start=1):
            set_cell(table.cell(r, c), get(model_name, cohort, param))

    for row in table.rows:
        for cell in row.cells:
            clear_all_borders(cell)
    apply_three_line_borders(table, header_row=0, last_row=n_rows - 1, n_cols=len(headers))

doc.save(OUTPUT)
print(f"\nSaved: {OUTPUT}")
