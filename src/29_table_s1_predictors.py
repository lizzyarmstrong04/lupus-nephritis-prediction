"""
Supplementary Table S1: Final predictor sets by cohort.

Three-line rule style (matching the main-manuscript table convention, not
the full-grid style used for the other appendix tables) per explicit spec:
rule above header, below header, and at foot only - no vertical rules, no
shading, no grid.

Values verified against source before writing this script:
  - EPV_MAX for the 5-year cohort is 15 (not floor(166/10)=16): the actual
    analysis script (src/5_year/04_feature_selection_5yr.py) hardcodes
    EPV_MAX=15 with comment "166 events / 10 = 16.6; cap at 15" - a
    deliberately conservative cap below the literal floor(), and that is
    what was actually enforced in the LASSO step and used consistently
    elsewhere (Methods S4.2 table, Table S4/S5). The floor() formula in
    the Methods prose is the imprecise part, not this value.
  - 5-year predictor count is 10 (not 11): Data/Processed/lupus_5yr_
    selected_clean.xlsx does still contain the stray dsDNA/SM/APL column
    on disk (confirming the data-integrity note in generate_methods_doc.py),
    but that column was explicitly dropped as a post-selection correction
    (Methods S4.4) and is not part of the reported final model.

All values below are stored in PREDICTOR_DATA so they can be updated in
one place; cross-checked against src/generate_methods_doc.py Section 4
tables at time of writing.

Saves: outputs/Table_S1_Predictors.docx
"""
from docx import Document
from docx.shared import Pt, Cm, Mm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

BASE = "/Users/elizabetharmstrong/Library/CloudStorage/OneDrive-ImperialCollegeLondon/Lupus Project"
OUTPUT = f"{BASE}/outputs/Table_S1_Predictors.docx"
FONT = "Times New Roman"

# ---------------------------------------------------------------------------
# All table content lives here so it can be updated in one place.
# ---------------------------------------------------------------------------
PREDICTOR_DATA = [
    {
        "cohort": "1-year flare",
        "n": 9, "epv_max": 9, "epv": 11.0,
        "predictors": "% chronic glomeruli; % glomeruli with necrosis; LN class; % active "
                       "glomeruli; % glomeruli with crescents; age at biopsy; ethnicity; "
                       "proteinuria at biopsy (log uPCR); C4 at biopsy",
    },
    {
        "cohort": "5-year flare",
        "n": 10, "epv_max": 15, "epv": 16.6,
        "predictors": "% chronic glomeruli; % glomeruli with necrosis; LN class; % active "
                       "glomeruli; % sclerosed glomeruli; age at biopsy; ethnicity; reason for "
                       "biopsy; CKD-EPI eGFR; prior cyclophosphamide exposure",
    },
    {
        "cohort": "Serial biopsy",
        "n": 2, "epv_max": 4, "epv": 17.0,
        "predictors": "Δ % chronic glomeruli (coefficient +0.291); time between biopsies, "
                       "months (coefficient −0.042)",
    },
    {
        "cohort": "ESRD 5-year",
        "n": 5, "epv_max": 11, "epv": 22.4,
        "predictors": "Creatinine at biopsy; %IFTA; CKD-EPI eGFR; % chronic glomeruli; "
                       "subepithelial deposit category",
    },
    {
        "cohort": "ESRD 10-year",
        "n": 17, "epv_max": 17, "epv": 10.3,
        "predictors": "Age at biopsy; %IFTA; creatinine at biopsy; % chronic glomeruli; "
                       "crescents; CKD-EPI eGFR; thrombotic microangiopathy; LN class; "
                       "C3 at biopsy; C4 low; prior cyclophosphamide; gender; subepithelial "
                       "deposit category; capillary wall IgM; number of globally sclerosed "
                       "glomeruli; biopsy number; number of glomeruli with crescents",
    },
]

CAPTION_LABEL = "Table S1 | "
CAPTION_TITLE = "Final predictor sets by cohort."

FOOTNOTE = (
    "EPV_MAX = ⌊events/10⌋, the hard ceiling imposed within the LASSO step; EPV actual = "
    "events divided by the number of predictors finally retained. The serial-biopsy model "
    "retained two predictors despite a cap of four, reflecting a discontinuity in the LASSO "
    "regularisation path rather than a deliberate constraint (Supplementary Methods S2). "
    "ESRD 10-year selected exactly at its cap; all other cohorts selected below it. eGFR "
    "calculated by CKD-EPI without ethnicity coefficient. ESRD, end-stage renal disease; "
    "EPV, events per variable; IFTA, interstitial fibrosis and tubular atrophy; uPCR, urine "
    "protein–creatinine ratio."
)

HEADERS = ["Cohort", "Predictors (n)", "EPV_MAX", "EPV actual", "Final predictor set"]
COL_WIDTHS_CM = [2.6, 1.6, 1.6, 1.6, 9.6]  # sums to 17.0 cm

CENTERED_COLS = {1, 2, 3}  # Predictors(n), EPV_MAX, EPV actual
LEFT_COLS = {0, 4}         # Cohort, Final predictor set


# ---------------------------------------------------------------------------
# Formatting helpers
# ---------------------------------------------------------------------------
def set_run_font(run, size=9, bold=False, italic=False):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
        rFonts.set(qn(attr), FONT)


def set_cell_text(cell, text, size=9, bold=False, italic=False,
                   align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    for r in list(p.runs):  # cell.text = "" can leave a stray unformatted empty run
        r._element.getparent().remove(r._element)
    p.alignment = align
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic)


def set_cell_margins(cell, top=40, bottom=40, left=100, right=100):
    """Cell margins in DXA (twentieths of a point)."""
    tcPr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement('w:tcMar')
    for tag, val in (('top', top), ('bottom', bottom), ('left', left), ('right', right)):
        node = OxmlElement(f'w:{tag}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        mar.append(node)
    tcPr.append(mar)


def set_cell_width_dxa(cell, width_cm):
    """Set an explicit cell width in DXA (1 cm = 566.929 dxa)."""
    dxa = int(round(width_cm * 566.929))
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = tcPr.find(qn('w:tcW'))
    if tcW is None:
        tcW = OxmlElement('w:tcW')
        tcPr.append(tcW)
    tcW.set(qn('w:w'), str(dxa))
    tcW.set(qn('w:type'), 'dxa')


def set_table_grid_widths_dxa(table, widths_cm):
    """Set the table-level <w:tblGrid> column widths in DXA."""
    tbl = table._tbl
    tblGrid = tbl.find(qn('w:tblGrid'))
    if tblGrid is None:
        tblGrid = OxmlElement('w:tblGrid')
        tbl.insert(0, tblGrid)
    for gridCol in list(tblGrid):
        tblGrid.remove(gridCol)
    for w in widths_cm:
        dxa = int(round(w * 566.929))
        gridCol = OxmlElement('w:gridCol')
        gridCol.set(qn('w:w'), str(dxa))
        tblGrid.append(gridCol)


def set_row_no_split(row):
    """Direct OOXML equivalent of 'do not allow this row to break across pages'."""
    trPr = row._tr.get_or_add_trPr()
    cantSplit = OxmlElement('w:cantSplit')
    trPr.append(cantSplit)


def set_cell_border(cell, edge, size=8, color="000000", val="single"):
    """Set a single border edge (top/bottom/left/right) on a cell via direct OOXML."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    tag = qn(f'w:{edge}')
    el = tcBorders.find(tag)
    if el is None:
        el = OxmlElement(f'w:{edge}')
        tcBorders.append(el)
    el.set(qn('w:val'), val)
    el.set(qn('w:sz'), str(size))
    el.set(qn('w:space'), '0')
    el.set(qn('w:color'), color)


def apply_three_line_borders(table, n_header_rows=1):
    """Rule above header, rule below header, rule at foot. No vertical rules, no grid.

    Clears every border on every cell first (the table carries no named style,
    but clearing explicitly guards against any inherited default), then adds
    back only the three horizontal rules.
    """
    n_rows = len(table.rows)
    n_cols = len(table.columns)
    header_row_idx = n_header_rows - 1
    last_row_idx = n_rows - 1
    for r in range(n_rows):
        for c in range(n_cols):
            cell = table.cell(r, c)
            for edge in ("top", "bottom", "left", "right"):
                set_cell_border(cell, edge, val="nil", size=0)
    for c in range(n_cols):
        set_cell_border(table.cell(0, c), "top", size=12)
        set_cell_border(table.cell(header_row_idx, c), "bottom", size=8)
        set_cell_border(table.cell(last_row_idx, c), "bottom", size=12)


# ---------------------------------------------------------------------------
# Build document
# ---------------------------------------------------------------------------
doc = Document()
section = doc.sections[0]
section.page_width = Mm(210)   # A4 portrait
section.page_height = Mm(297)
section.top_margin = section.bottom_margin = Cm(2)
section.left_margin = section.right_margin = Cm(2)

# Caption above the table
caption = doc.add_paragraph()
run = caption.add_run(CAPTION_LABEL)
set_run_font(run, size=10, bold=True)
run = caption.add_run(CAPTION_TITLE)
set_run_font(run, size=10, bold=False)
caption.paragraph_format.space_after = Pt(6)

# Table
table = doc.add_table(rows=1 + len(PREDICTOR_DATA), cols=len(HEADERS))
# No named table style (avoids inheriting default grid lines/shading) - all
# borders are cleared and set explicitly in apply_three_line_borders() below.
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False

set_table_grid_widths_dxa(table, COL_WIDTHS_CM)

# Header row
for c, h in enumerate(HEADERS):
    cell = table.cell(0, c)
    align = WD_ALIGN_PARAGRAPH.CENTER if c in CENTERED_COLS else WD_ALIGN_PARAGRAPH.LEFT
    set_cell_text(cell, h, size=9, bold=True, align=align)
    set_cell_width_dxa(cell, COL_WIDTHS_CM[c])
    set_cell_margins(cell)

# Data rows
for r, row_data in enumerate(PREDICTOR_DATA, start=1):
    values = [
        row_data["cohort"],
        str(row_data["n"]),
        str(row_data["epv_max"]),
        f'{row_data["epv"]:.1f}',
        row_data["predictors"],
    ]
    for c, val in enumerate(values):
        cell = table.cell(r, c)
        align = WD_ALIGN_PARAGRAPH.CENTER if c in CENTERED_COLS else WD_ALIGN_PARAGRAPH.LEFT
        set_cell_text(cell, val, size=9, align=align)
        set_cell_width_dxa(cell, COL_WIDTHS_CM[c])
        set_cell_margins(cell)

# No vertical rules / shading anywhere; three-line borders only
apply_three_line_borders(table, n_header_rows=1)

# Prevent rows breaking across a page boundary
for row in table.rows:
    set_row_no_split(row)

# Footnote
footnote = doc.add_paragraph()
run = footnote.add_run(FOOTNOTE)
set_run_font(run, size=8, italic=True)
footnote.paragraph_format.space_before = Pt(6)

doc.save(OUTPUT)
print(f"Saved: {OUTPUT}")
