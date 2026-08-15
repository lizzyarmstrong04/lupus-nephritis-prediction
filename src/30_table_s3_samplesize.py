"""
Table S3 (Appendix): Post-hoc minimum sample size by cohort (pmsampsize).

Condensed companion to Table S5 (src/27_table_s5_samplesize.py) - same
underlying data (outputs/pmsampsize_results.xlsx, from
src/22_pmsampsize.py), fewer columns: Cohort, p, Prevalence, C-stat
(LR, BC), n Required, n Actual, EPV Actual, Adequate? (drops Cox-Snell R2,
Events Required, EPP Required, Events Actual - all present in Table S5 for
readers who want the full derivation).

Every value below is read directly from outputs/pmsampsize_results.xlsx,
not retyped - and was independently cross-checked against the user-supplied
target table before writing this script (exact match on all 5 cohorts x
6 numeric columns).

Three-line rule style (rule above header, below header, at foot; no
vertical rules, no grid), matching the main-manuscript Table 1/2 convention
(src/25_tables_1_2.py) and Tables S1/S4/S5/S6, Times New Roman.

Saves: outputs/Table_S3_SampleSize_Condensed.docx
"""
import pandas as pd
from docx import Document
from docx.shared import Pt, Cm, Mm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

BASE = "/Users/elizabetharmstrong/Library/CloudStorage/OneDrive-ImperialCollegeLondon/Lupus Project"
OUT = f"{BASE}/outputs"
OUTPUT = f"{OUT}/Table_S3_SampleSize_Condensed.docx"
FONT = "Times New Roman"

df = pd.read_excel(f"{OUT}/pmsampsize_results.xlsx")
print(df.to_string(index=False))

# --- Cross-check against the user-supplied target table before building ---
TARGET = {
    "1-Year flare":  {"p": 9,  "prev": 0.230, "cstat": 0.709, "n_req": 795, "n_act": 430, "epv": 11.00, "adeq": "No"},
    "5-Year flare":  {"p": 10, "prev": 0.466, "cstat": 0.670, "n_req": 972, "n_act": 356, "epv": 16.60, "adeq": "No"},
    "Serial biopsy": {"p": 2,  "prev": 0.486, "cstat": 0.660, "n_req": 384, "n_act": 70,  "epv": 17.00, "adeq": "No"},
    "ESRD 5-Year":   {"p": 5,  "prev": 0.141, "cstat": 0.797, "n_req": 293, "n_act": 796, "epv": 22.40, "adeq": "Yes"},
    "ESRD 10-Year":  {"p": 17, "prev": 0.220, "cstat": 0.817, "n_req": 621, "n_act": 796, "epv": 10.29, "adeq": "Yes"},
}
NAME_MAP = {"1-Year flare": "1-Year flare", "5-Year flare": "5-Year flare",
            "Serial biopsy": "Serial biopsy", "ESRD 5-Year": "ESRD 5-Year",
            "ESRD 10-Year": "ESRD 10-Year"}
for row in df.itertuples(index=False):
    cohort = str(row.Cohort)
    t = TARGET.get(cohort)
    if t is None:
        continue
    checks = [
        (row._1 == t["p"], "p"),
        (abs(row.Prevalence - t["prev"]) < 1e-6, "Prevalence"),
        (abs(row._3 - t["cstat"]) < 1e-6, "C-stat"),
        (row._5 == t["n_req"], "n Required"),
        (row._8 == t["n_act"], "n Actual"),
        (abs(row._10 - t["epv"]) < 1e-6, "EPV Actual"),
        (str(row._11) == t["adeq"], "Adequate?"),
    ]
    for ok, label in checks:
        assert ok, f"MISMATCH for {cohort} / {label}"
print("\nAll values verified against target table - exact match.\n")


# --- Border / font helpers (matching src/25_tables_1_2.py exactly) ---

def set_cell_borders(cell, top=None, bottom=None, left=None, right=None, insideH=None, insideV=None):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    specs = {"top": top, "bottom": bottom, "left": left, "right": right,
             "insideH": insideH, "insideV": insideV}
    for edge, spec in specs.items():
        el = tcBorders.find(qn(f'w:{edge}'))
        if el is None:
            el = OxmlElement(f'w:{edge}')
            tcBorders.append(el)
        if spec is None:
            el.set(qn('w:val'), 'nil')
        else:
            el.set(qn('w:val'), 'single')
            el.set(qn('w:sz'), str(spec))
            el.set(qn('w:color'), '000000')
            el.set(qn('w:space'), '0')


def clear_all_borders(cell):
    set_cell_borders(cell)


def apply_three_line_borders(table, header_row, last_row, n_cols):
    THICK, THIN = 18, 8
    if header_row == 0:
        for c in range(n_cols):
            set_cell_borders(table.cell(0, c), top=THICK, bottom=THIN)
    else:
        for c in range(n_cols):
            set_cell_borders(table.cell(0, c), top=THICK)
        for c in range(n_cols):
            set_cell_borders(table.cell(header_row, c), bottom=THIN)
    for c in range(n_cols):
        set_cell_borders(table.cell(last_row, c), bottom=THICK)


def set_table_cell_margins(table, top_pt=3, bottom_pt=3, left_pt=5, right_pt=5):
    tbl = table._tbl
    tblPr = tbl.tblPr
    cellMar = OxmlElement('w:tblCellMar')
    for edge, pt_val in (('top', top_pt), ('bottom', bottom_pt), ('left', left_pt), ('right', right_pt)):
        node = OxmlElement(f'w:{edge}')
        node.set(qn('w:w'), str(int(pt_val * 20)))
        node.set(qn('w:type'), 'dxa')
        cellMar.append(node)
    tblPr.append(cellMar)


def set_run_font(run, size=10, bold=False, italic=False):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
        rFonts.set(qn(attr), FONT)


def set_cell(cell, text, size=10, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ""
    p = cell.paragraphs[0]
    for r in list(p.runs):  # cell.text = "" can leave a stray unformatted empty run
        r._element.getparent().remove(r._element)
    p.alignment = align
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic)


doc = Document()
section = doc.sections[0]
section.page_width = Mm(210)   # A4 portrait - 8 columns fits without landscape
section.page_height = Mm(297)
section.top_margin = section.bottom_margin = Cm(2)
section.left_margin = section.right_margin = Cm(2)

title = doc.add_paragraph()
run = title.add_run("Table S3. Post-hoc minimum sample size by cohort (pmsampsize).")
set_run_font(run, size=11, bold=True)
title.paragraph_format.space_after = Pt(8)

headers = ["Cohort", "p", "Prevalence", "C-stat\n(LR, BC)", "n\nRequired",
           "n\nActual", "EPV\nActual", "Adequate?"]
col_widths = [2.6, 1.2, 1.8, 2.0, 1.8, 1.6, 1.6, 1.8]  # sums to 14.4cm, fits 17cm usable

n_rows = 1 + len(df)
table = doc.add_table(rows=n_rows, cols=len(headers))
table.style = "Table Grid"
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False
set_table_cell_margins(table, top_pt=3, bottom_pt=3, left_pt=5, right_pt=5)

for i, w in enumerate(col_widths):
    table.columns[i].width = Cm(w)
for row in table.rows:
    for cell, w in zip(row.cells, col_widths):
        cell.width = Cm(w)

for c, h in enumerate(headers):
    set_cell(table.cell(0, c), h, size=9.5, bold=True)

for r, row in enumerate(df.itertuples(index=False), start=1):
    adequate = row._11  # "Adequate per pmsampsize?"
    set_cell(table.cell(r, 0), row.Cohort, align=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell(table.cell(r, 1), str(row._1))                          # Parameters (p)
    set_cell(table.cell(r, 2), f"{row.Prevalence:.3f}")              # Prevalence
    set_cell(table.cell(r, 3), f"{row._3:.3f}")                      # Anticipated C-statistic
    set_cell(table.cell(r, 4), str(row._5))                          # n required
    set_cell(table.cell(r, 5), str(row._8))                          # n actual
    set_cell(table.cell(r, 6), f"{row._10:.2f}")                     # EPV actual
    set_cell(table.cell(r, 7), adequate, bold=True)                  # Adequate?

for row in table.rows:
    for cell in row.cells:
        clear_all_borders(cell)
apply_three_line_borders(table, header_row=0, last_row=n_rows - 1, n_cols=len(headers))

footnote = doc.add_paragraph()
run = footnote.add_run(
    "p = number of final predictor parameters. Anticipated C-statistic = the cohort's own "
    "optimism-adjusted (bias-corrected, Harrell bootstrap) Logistic Regression AUROC. Required n "
    "per Riley et al. 2020 (BMJ 368:m441): max of predictor-effect shrinkage (≤10% loss), "
    "Nagelkerke R² difference (≤0.05), and intercept precision (0.05 margin) criteria. See "
    "Table S5 for the full per-criterion derivation (Cox-Snell R², events required, EPP required, "
    "events actual)."
)
set_run_font(run, size=9, italic=True)
footnote.paragraph_format.space_before = Pt(6)

doc.save(OUTPUT)
print(f"\nSaved: {OUTPUT}")
